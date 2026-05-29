"""Peek at Q2.1 (which has fraction lines) and the intro section."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
from pathlib import Path
from docx import Document
SRC = Path(r"school-agents/outputs/homework/sci30-unitc-emr-review-v1/sci30-unitc-emr-review-student.docx")
doc = Document(str(SRC))

print("===== Intro section (paragraphs 0-19) =====")
for i in range(20):
    if i >= len(doc.paragraphs):
        break
    p = doc.paragraphs[i]
    has_math = any("oMath" in (c.tag or "") for c in p._p.iterchildren())
    sub_runs = [(r.text or "", bool(r.font.subscript)) for r in p.runs]
    n_sub = sum(1 for _, s in sub_runs if s)
    print(f"  {i:3d} math={int(has_math)} subs={n_sub} text={p.text[:90]!r}")

print("\n===== Q2.1 ====")
start = None
for i, p in enumerate(doc.paragraphs):
    if "Question 2.1" in p.text:
        start = i
        break
print(f"Q2.1 starts at paragraph {start}")

for j in range(0, 50):
    if start + j >= len(doc.paragraphs):
        break
    p = doc.paragraphs[start + j]
    has_math = any("oMath" in (c.tag or "") for c in p._p.iterchildren())
    sub_runs = [(r.text or "", bool(r.font.subscript)) for r in p.runs]
    sub_runs_short = [(t[:50], s) for t, s in sub_runs if t]
    print(f"  +{j:3d} math={int(has_math)} runs={sub_runs_short[:5]}")
