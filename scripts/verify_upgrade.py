"""Verify the upgraded DOCX preserves columns/page-breaks and contains
proper OMML math zones."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

SRC = Path(r"school-agents/outputs/homework/sci30-unitc-emr-review-v1/sci30-unitc-emr-review-student.docx")

doc = Document(str(SRC))

print(f"Sections: {len(doc.sections)}")
for i, sec in enumerate(doc.sections):
    sectPr = sec._sectPr
    cols = sectPr.find(qn("w:cols"))
    cols_attrs = dict(cols.attrib) if cols is not None else None
    print(f"  Section {i}: cols={cols_attrs}")

print(f"\nParagraphs: {len(doc.paragraphs)}")

# Count OMML elements
omath_count = 0
for p in doc.paragraphs:
    for child in p._p.iterchildren():
        if "oMath" in (child.tag or ""):
            omath_count += 1
print(f"OMML <m:oMath> elements found: {omath_count}")

# Count subscript runs
sub_runs = 0
for p in doc.paragraphs:
    for r in p.runs:
        if r.font.subscript:
            sub_runs += 1
print(f"Run-level subscript runs: {sub_runs}")

# Page breaks (counted by raw <w:br w:type="page"/>)
import lxml.etree as ET
xml = doc.element.body.xml
pb_count = xml.count('<w:br w:type="page"')
print(f"Page-break elements: {pb_count}")

# Sample first 3 question paragraphs that contain math
print("\nSample paragraphs near Q1.1 and Q2.1:")
for i, p in enumerate(doc.paragraphs):
    text = p.text
    if "Question 1.1" in text or "Question 2.1" in text:
        print(f"\n--- {text} ---")
        for j in range(8):
            pp = doc.paragraphs[i + j]
            ptxt = pp.text[:120].replace("\n", " | ")
            has_math = any("oMath" in (c.tag or "") for c in pp._p.iterchildren())
            sub_count = sum(1 for r in pp.runs if r.font.subscript)
            print(f"  +{j}  math={int(has_math)} subs={sub_count}  {ptxt}")
