"""Peek at Q1.1 GRAS(S) content to confirm math zones inserted."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
from pathlib import Path
from docx import Document
SRC = Path(r"school-agents/outputs/homework/sci30-unitc-emr-review-v1/sci30-unitc-emr-review-student.docx")
doc = Document(str(SRC))

# Find Q1.1 prompt and step through 30 paragraphs
start = None
for i, p in enumerate(doc.paragraphs):
    if "Question 1.1" in p.text:
        start = i
        break
print(f"Q1.1 starts at paragraph {start}")

for j in range(0, 50):
    if start + j >= len(doc.paragraphs):
        break
    p = doc.paragraphs[start + j]
    has_math = any("oMath" in (c.tag or "") for c in p._p.iterchildren())
    sub_runs = [(r.text or "", bool(r.font.subscript)) for r in p.runs]
    sub_runs_short = [(t[:40], s) for t, s in sub_runs if t]
    print(f"+{j:3d} math={int(has_math)} runs={sub_runs_short[:6]}{'...' if len(sub_runs_short)>6 else ''}")
