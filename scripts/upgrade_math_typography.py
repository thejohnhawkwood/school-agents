"""
Upgrade the math typography in the Sci 30 Unit C review DOCX in place,
preserving the user's column / page-break formatting.

Pipeline:
  Pass 1: known equation paragraphs (e.g. "F⃗g = mg⃗", "g = Gm / r²", etc.)
          and the section "Equation(s) used" lines are converted to OMML
          math zones using a curated LaTeX lookup table.

  Pass 2: every indented "Substitute and solve" line that is pure math
          (no English prose words) is converted to an OMML math zone with:
            - compound units (N/kg, m/s, kW·h, ...) protected as upright text
            - simple units after numbers (kg, m, s, V, A, J, Hz, Ω, ...) also
              wrapped as upright text
            - all remaining `/` between atoms rendered as built-up fractions

  Pass 3: any remaining prose paragraph that contains an underscore
          identifier (R_T, M_Earth, m_total, F_g, ...) or a vector-arrow
          identifier (F⃗g, F⃗g,Mars, ...) gets run-level subscripts so the
          subscript visibly drops below the baseline in Word.

We use latex2mathml + a custom MathML->OMML converter (mathml_to_omml.py).
The custom converter avoids two bugs in mathml2omml v0.0.2:
  - malformed <m:groupChrPr> in vector accents
  - unnecessary <m:box> wrappers (Word renders these as visible empty
    boxes, which the user reported as "□(→┬□F_□g...").
"""

from __future__ import annotations
import io
import re
import sys
from copy import deepcopy
from pathlib import Path

# Ensure stdout can carry the unicode in our debug prints (vector arrows etc.)
if hasattr(sys.stdout, "buffer"):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from lxml import etree

import latex2mathml.converter as l2m

sys.path.insert(0, str(Path(__file__).parent))
import mathml_to_omml as m2o   # noqa: E402

DOCX = Path(r"school-agents/outputs/homework/sci30-unitc-emr-review-v1/sci30-unitc-emr-review-student.docx")

OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


# =====================================================================================
# Equation lookup -- known paragraph text -> LaTeX
# =====================================================================================


EQUATION_LATEX: dict[str, str] = {
    # Type 1
    "F\u20D7g = mg\u20D7": r"\vec{F}_{g} = m\vec{g}",

    # Type 2
    "g = Gm / r\u00B2": r"g = \frac{Gm}{r^{2}}",

    # Type 3
    "|E\u20D7| = kq / r\u00B2": r"\left|\vec{E}\right| = \frac{kq}{r^{2}}",
    "rearranged: r = \u221A(kq / |E\u20D7|)": r"r = \sqrt{\frac{kq}{|\vec{E}|}}",

    # Type 4
    "V = IR": r"V = IR",
    "V = IR  \u2192  R = V/I,  I = V/R":
        r"V = IR \;\Rightarrow\; R = \frac{V}{I},\;\; I = \frac{V}{R}",
    "V = IR  \u2192  I = V/R": r"V = IR \;\Rightarrow\; I = \frac{V}{R}",
    "V = IR (across each component)": r"V = IR",
    "V = IR (each branch has full source voltage)": r"V = IR",
    "V = IR  \u2192  I_total = V / R_T":
        r"V = IR \;\Rightarrow\; I_{\text{total}} = \frac{V}{R_{T}}",
    "I_total = V / R_T": r"I_{\text{total}} = \frac{V}{R_{T}}",

    # Type 5 - series
    "R_T = R\u2081 + R\u2082 + ... + R_n":
        r"R_{T} = R_{1} + R_{2} + \ldots + R_{n}",
    "R_T = R\u2081 + R\u2082 + R\u2083 + R\u2084":
        r"R_{T} = R_{1} + R_{2} + R_{3} + R_{4}",
    "R_T = R\u2081 + R\u2082 + ... + R\u2088":
        r"R_{T} = R_{1} + R_{2} + \ldots + R_{8}",
    "R_T = R\u2081 + R\u2082 + R\u2083":
        r"R_{T} = R_{1} + R_{2} + R_{3}",

    # Type 6 - parallel
    "1/R_T = 1/R\u2081 + 1/R\u2082 + ... + 1/R_n":
        r"\frac{1}{R_{T}} = \frac{1}{R_{1}} + \frac{1}{R_{2}} + \ldots + \frac{1}{R_{n}}",
    "1/R_T = 1/R\u2081 + 1/R\u2082":
        r"\frac{1}{R_{T}} = \frac{1}{R_{1}} + \frac{1}{R_{2}}",
    "1/R_T = 1/R\u2081 + 1/R\u2082 + 1/R\u2083":
        r"\frac{1}{R_{T}} = \frac{1}{R_{1}} + \frac{1}{R_{2}} + \frac{1}{R_{3}}",
    "1/R_T = n / R_each   (when all R are equal)":
        r"\frac{1}{R_{T}} = \frac{n}{R_{\text{each}}}",

    # Type 7 - power
    "P = IV": r"P = IV",
    "P = I\u00B2R": r"P = I^{2}R",
    "P = IV (for the verification)": r"P = IV",

    # Type 8 - energy
    "E = Pt": r"E = Pt",
    "1 kW\u00B7h = 3.60 \u00D7 10\u2076 J":
        r"1\,\text{kW}\!\cdot\!\text{h} = 3.60 \times 10^{6}\,\text{J}",
    "Cost = E \u00D7 rate":
        r"\text{Cost} = E \times \text{rate}",

    # Type 9 - transformers
    "N_p / N_s = V_p / V_s":
        r"\frac{N_{p}}{N_{s}} = \frac{V_{p}}{V_{s}}",
    "N_p / N_s = I_s / I_p":
        r"\frac{N_{p}}{N_{s}} = \frac{I_{s}}{I_{p}}",
    "V_p / V_s = I_s / I_p":
        r"\frac{V_{p}}{V_{s}} = \frac{I_{s}}{I_{p}}",
    "N_p / N_s = V_p / V_s  \u2192  N_s = N_p \u00D7 (V_s / V_p)":
        r"\frac{N_{p}}{N_{s}} = \frac{V_{p}}{V_{s}} \;\Rightarrow\; "
        r"N_{s} = N_{p}\cdot\frac{V_{s}}{V_{p}}",
    "I_s / I_p = V_p / V_s  \u2192  I_s = I_p \u00D7 (V_p / V_s)":
        r"\frac{I_{s}}{I_{p}} = \frac{V_{p}}{V_{s}} \;\Rightarrow\; "
        r"I_{s} = I_{p}\cdot\frac{V_{p}}{V_{s}}",

    # Type 10 - waves
    "v = f\u03BB": r"v = f\lambda",
    "c = f\u03BB": r"c = f\lambda",
    "c = 3.00 \u00D7 10\u2078 m/s":
        r"c = 3.00 \times 10^{8}\,\text{m/s}",
    "c = f\u03BB  \u2192  \u03BB = c / f":
        r"c = f\lambda \;\Rightarrow\; \lambda = \frac{c}{f}",
    "c = f\u03BB  \u2192  f = c / \u03BB":
        r"c = f\lambda \;\Rightarrow\; f = \frac{c}{\lambda}",
}


# =====================================================================================
# Unit handling -- compound units (with "/") MUST be matched first
# =====================================================================================


COMPOUND_UNITS = [
    "N\u00B7m\u00B2/kg\u00B2",     # N·m²/kg²  (G)
    "N\u00B7m\u00B2/C\u00B2",      # N·m²/C²   (k)
    "kg\u00B7m/s\u00B2",            # kg·m/s²
    "kW\u00B7h",                    # kW·h
    "MW\u00B7h",
    "N/kg",
    "N/C",
    "m/s\u00B2",
    "m/s",
    "kg/m\u00B3",
    "J/s",
    "J/kg",
    "J/(kg\u00B7K)",
    "rad/s",
    "1/s",
]
COMPOUND_UNITS.sort(key=len, reverse=True)

SIMPLE_UNITS = {
    "kg", "g", "mg",
    "m", "cm", "mm", "km", "nm", "\u03BCm",
    "s", "ms", "min", "h", "hr",
    "Hz", "kHz", "MHz", "GHz",
    "J", "kJ", "MJ", "GJ", "eV", "keV", "MeV",
    "W", "kW", "MW",
    "V", "mV", "kV",
    "A", "mA", "\u03BCA",
    "C", "mC", "\u03BCC", "nC", "pC",
    "\u03A9", "k\u03A9", "M\u03A9",
    "T", "Wb",
    "K", "Pa", "kPa", "MPa", "GPa",
    "rad", "sr", "mol",
    "N",
    "L", "mL",
}


# =====================================================================================
# Subscript regexes (Pass 3)
# =====================================================================================


UNDER_SUB = re.compile(
    r"(?<![A-Za-z0-9])"
    r"([A-Za-z])"
    r"_"
    r"([A-Za-z0-9][A-Za-z0-9]*"
    r"(?:,[A-Za-z][A-Za-z0-9]*)*)"
)
VEC_SUB = re.compile(
    r"(?<![A-Za-z0-9])"
    r"([A-Z])\u20D7"
    r"([a-z][A-Za-z0-9]*"
    r"(?:,[A-Za-z][A-Za-z0-9]*)*)"
)


# =====================================================================================
# OMML helpers
# =====================================================================================


def latex_to_omml_element(latex: str) -> etree._Element:
    """Convert a LaTeX string to an <m:oMath> XML element."""
    mathml = l2m.convert(latex)
    omml_str = m2o.convert(mathml)
    elem = etree.fromstring(omml_str)
    return elem


def remove_runs_and_math(p) -> None:
    """Remove all w:r and m:oMath children of the paragraph; keep w:pPr."""
    for child in list(p._p.iterchildren()):
        local = etree.QName(child.tag).localname
        if local in ("r", "oMath"):
            p._p.remove(child)


def append_text_run(p, text: str, *, subscript: bool = False) -> None:
    if not text:
        return
    r = p.add_run(text)
    if subscript:
        r.font.subscript = True


def append_omml(p, omml_elem) -> None:
    p._p.append(deepcopy(omml_elem))


def replace_with_math_zone(p, latex: str, *,
                           prefix_text: str = "",
                           suffix_text: str = "") -> None:
    """Replace paragraph content with optional plain-text prefix, the math
    zone, and optional plain-text suffix."""
    omml = latex_to_omml_element(latex)
    remove_runs_and_math(p)
    if prefix_text:
        append_text_run(p, prefix_text)
    append_omml(p, omml)
    if suffix_text:
        append_text_run(p, suffix_text)


# =====================================================================================
# Unicode -> LaTeX preprocessing
# =====================================================================================


_SUP_MAP = {
    "\u00B2": "2", "\u00B3": "3", "\u2070": "0",
    "\u00B9": "1", "\u2074": "4", "\u2075": "5",
    "\u2076": "6", "\u2077": "7", "\u2078": "8",
    "\u2079": "9", "\u207B": "-", "\u207A": "+",
}
_SUB_MAP = {
    "\u2080": "0", "\u2081": "1", "\u2082": "2",
    "\u2083": "3", "\u2084": "4", "\u2085": "5",
    "\u2086": "6", "\u2087": "7", "\u2088": "8",
    "\u2089": "9",
}


def _fix_unicode_supersubs(s: str) -> str:
    """Convert consecutive Unicode super/sub digits to LaTeX ^{...} / _{...}."""
    out: list[str] = []
    i = 0
    while i < len(s):
        ch = s[i]
        if ch in _SUP_MAP:
            j = i
            buf = ""
            while j < len(s) and s[j] in _SUP_MAP:
                buf += _SUP_MAP[s[j]]
                j += 1
            out.append("^{" + buf + "}")
            i = j
        elif ch in _SUB_MAP:
            j = i
            buf = ""
            while j < len(s) and s[j] in _SUB_MAP:
                buf += _SUB_MAP[s[j]]
                j += 1
            out.append("_{" + buf + "}")
            i = j
        else:
            out.append(ch)
            i += 1
    return "".join(out)


_UNI2LATEX_SIMPLE = [
    ("\u00D7", r" \times "),
    ("\u00B7", r" \cdot "),
    ("\u00F7", r" \div "),
    ("\u2248", r" \approx "),
    ("\u2260", r" \neq "),
    ("\u2264", r" \le "),
    ("\u2265", r" \ge "),
    ("\u2192", r" \Rightarrow "),
    ("\u2026", r" \ldots "),
    ("\u00B0", r"^{\circ}"),
    ("\u03A9", r"\Omega "),
    ("\u03BC", r"\mu "),
    ("\u03BB", r"\lambda "),
    ("\u0394", r"\Delta "),
    ("\u03C0", r"\pi "),
    ("\u20D7", ""),    # vector accent (handled by VEC_PREFIX below)
    ("\u2261", r" \equiv "),
    ("\u2212", "-"),
    ("\u00B1", r" \pm "),
]


# =====================================================================================
# Vector pre-processing: F⃗ -> \vec{F}, F⃗_{T} -> \vec{F}_{T}
# =====================================================================================


def _replace_vectors(s: str) -> str:
    """Replace 'X⃗' (capital + combining arrow) with '\\vec{X}'.
    Also handle 'X⃗_letters' or 'X⃗letters' (no underscore) -- the
    subscript is everything after the arrow until a non-identifier char."""

    out: list[str] = []
    i = 0
    while i < len(s):
        if (i + 1 < len(s)
                and s[i].isalpha() and s[i].isupper()
                and s[i+1] == "\u20D7"):
            base = s[i]
            j = i + 2
            # Optional subscript letters following directly (no _) like F⃗g
            # or via underscore F⃗_g. We assume any letters/digits/comma
            # before whitespace/operator are subscript content.
            sub_chars: list[str] = []
            if j < len(s) and s[j] == "_":
                j += 1
                # bracketed: _{...}
                if j < len(s) and s[j] == "{":
                    depth = 1
                    j += 1
                    while j < len(s) and depth > 0:
                        if s[j] == "{":
                            depth += 1
                        elif s[j] == "}":
                            depth -= 1
                            if depth == 0:
                                j += 1
                                break
                        sub_chars.append(s[j])
                        j += 1
                else:
                    while j < len(s) and (s[j].isalnum() or s[j] == ","):
                        sub_chars.append(s[j])
                        j += 1
            else:
                # Greedy lowercase letters / commas / digits as subscript
                while j < len(s) and (s[j].islower() or s[j].isdigit() or s[j] == ","):
                    sub_chars.append(s[j])
                    j += 1

            if sub_chars:
                out.append(r"\vec{" + base + r"}_{" + "".join(sub_chars) + r"}")
            else:
                out.append(r"\vec{" + base + r"}")
            i = j
            continue
        # Lowercase + arrow: e.g. g⃗ -> \vec{g}
        if (i + 1 < len(s)
                and s[i].isalpha() and s[i].islower()
                and s[i+1] == "\u20D7"):
            out.append(r"\vec{" + s[i] + r"}")
            i += 2
            continue
        out.append(s[i])
        i += 1
    return "".join(out)


# =====================================================================================
# Unit protection / wrapping in \text{...}
# =====================================================================================


_UNIT_PLACEHOLDER = "\u0001U{}\u0001"


def _protect_compound_units(s: str) -> tuple[str, list[str]]:
    units: list[str] = []
    for u in COMPOUND_UNITS:
        while u in s:
            units.append(u)
            ph = _UNIT_PLACEHOLDER.format(len(units) - 1)
            s = s.replace(u, ph, 1)
    return s, units


def _restore_units_to_latex(s: str, units: list[str]) -> str:
    for idx, u in enumerate(units):
        ph = _UNIT_PLACEHOLDER.format(idx)
        s = s.replace(ph, r"\,\text{" + u + r"}")
    return s


def _wrap_simple_units_after_numbers(s: str) -> str:
    """Wrap simple units (kg, m, V, ...) following a number with optional
    space, in \\text{...}. Operates on the LaTeX-ish intermediate form."""
    units_alt = "|".join(re.escape(u) for u in sorted(SIMPLE_UNITS, key=len, reverse=True))
    # Number boundary: digit, ), }, or end of ^{...}
    pattern = re.compile(
        r"(\d|\}|\))(?:\s+|(?=\\,))?(" + units_alt + r")(?![A-Za-z])"
    )
    while True:
        new = pattern.sub(lambda m: m.group(1) + r"\,\text{" + m.group(2) + "}", s)
        if new == s:
            break
        s = new
    return s


# =====================================================================================
# Slash-to-fraction conversion using a small tokenizer
# =====================================================================================


def _tokenize_for_fractions(s: str):
    """Tokenize a LaTeX-ish string into (kind, text) tokens.
    Kinds: 'ATOM', 'SLASH', 'OP', 'WS' (skipped)."""
    tokens: list[tuple[str, str]] = []
    i = 0
    n = len(s)

    def read_braced(j: int) -> int:
        """If s[j] == '{', return position after the matching '}'; else j."""
        if j < n and s[j] == "{":
            depth = 1
            j += 1
            while j < n and depth > 0:
                if s[j] == "{":
                    depth += 1
                elif s[j] == "}":
                    depth -= 1
                    if depth == 0:
                        j += 1
                        break
                j += 1
        return j

    def read_atom_tail(j: int) -> int:
        """Allow trailing _{...}, ^{...} on an atom."""
        while j < n and s[j] in ("_", "^"):
            j += 1
            if j < n and s[j] == "{":
                j = read_braced(j)
            else:
                j += 1
        return j

    while i < n:
        c = s[i]
        if c.isspace():
            i += 1
            continue
        if c == "/":
            tokens.append(("SLASH", "/"))
            i += 1
            continue
        if c in "+=":
            tokens.append(("OP", c))
            i += 1
            continue
        if c == "-":
            tokens.append(("OP", c))
            i += 1
            continue
        if c == "\\":
            # LaTeX command \name with optional {...} arguments
            j = i + 1
            # Punctuation commands like \, \! \;
            if j < n and not s[j].isalpha():
                j += 1
                tokens.append(("ATOM", s[i:j]))
                i = j
                continue
            while j < n and s[j].isalpha():
                j += 1
            while j < n and s[j] == "{":
                j = read_braced(j)
            j = read_atom_tail(j)
            tokens.append(("ATOM", s[i:j]))
            i = j
            continue
        if c == "(":
            depth = 1
            j = i + 1
            while j < n and depth > 0:
                if s[j] == "(":
                    depth += 1
                elif s[j] == ")":
                    depth -= 1
                j += 1
            j = read_atom_tail(j)
            tokens.append(("ATOM", s[i:j]))
            i = j
            continue
        if c == "|":
            j = i + 1
            while j < n and s[j] != "|":
                j += 1
            if j < n:
                j += 1
            tokens.append(("ATOM", s[i:j]))
            i = j
            continue
        # Identifier / number / placeholder
        if c.isalnum() or c == "." or c == "\u0001":
            j = i
            while j < n and (s[j].isalnum() or s[j] == "." or s[j] == "\u0001"
                             or s[j] in "_^"):
                if s[j] in ("_", "^"):
                    j += 1
                    if j < n and s[j] == "{":
                        j = read_braced(j)
                    else:
                        j += 1
                else:
                    j += 1
            tokens.append(("ATOM", s[i:j]))
            i = j
            continue
        # Other char (unicode, comma, etc.)
        tokens.append(("OP", c))
        i += 1
    return tokens


def _strip_outer_parens(s: str) -> str:
    """If s == "(...)" with balanced parens at outermost, return inner."""
    s = s.strip()
    if not (s.startswith("(") and s.endswith(")")):
        return s
    depth = 0
    for i, c in enumerate(s):
        if c == "(":
            depth += 1
        elif c == ")":
            depth -= 1
            if depth == 0 and i != len(s) - 1:
                return s
    return s[1:-1]


def _consume_fractions(tokens):
    """Replace any [ATOM, SLASH, ATOM] sequence with a single [\\frac] ATOM."""
    out: list[tuple[str, str]] = []
    i = 0
    while i < len(tokens):
        if (i + 2 < len(tokens)
                and tokens[i][0] == "ATOM"
                and tokens[i + 1][0] == "SLASH"
                and tokens[i + 2][0] == "ATOM"):
            num = _strip_outer_parens(tokens[i][1])
            den = _strip_outer_parens(tokens[i + 2][1])
            out.append(("ATOM", r"\frac{" + num + r"}{" + den + r"}"))
            i += 3
        else:
            out.append(tokens[i])
            i += 1
    return out


def _render_tokens(tokens) -> str:
    parts = []
    for kind, text in tokens:
        if kind == "OP":
            parts.append(" " + text + " ")
        else:
            parts.append(text)
    return "".join(parts)


def expression_to_latex(s: str) -> str:
    """Take a math expression string in our internal text form and produce
    a LaTeX string with units in \\text{} and all `/` converted to fractions."""
    s = _replace_vectors(s)

    # Escape characters that LaTeX would otherwise treat specially.
    s = s.replace("%", r"\%")

    # Protect compound units before any other transformation (they may
    # contain "/" which we don't want to fracture).
    s, compound = _protect_compound_units(s)

    # Unicode -> LaTeX
    s = _fix_unicode_supersubs(s)
    for u, lat in _UNI2LATEX_SIMPLE:
        s = s.replace(u, lat)

    # Variable subscript: var_word -> var_{word}
    s = re.sub(r"([A-Za-z])_([A-Za-z][A-Za-z0-9]*)", r"\1_{\2}", s)

    # Wrap simple units after numbers in \text{...}
    s = _wrap_simple_units_after_numbers(s)

    # Tokenize and convert / to fractions (compound units are still
    # placeholders, so / inside them is not visible)
    tokens = _tokenize_for_fractions(s)
    tokens = _consume_fractions(tokens)
    s = _render_tokens(tokens)

    # Restore compound units as upright text
    s = _restore_units_to_latex(s, compound)

    # Tidy up multiple spaces
    s = re.sub(r"\s+", " ", s).strip()
    return s


# =====================================================================================
# Pass 1 + Pass 2 detection
# =====================================================================================


_PROSE_WORDS = re.compile(
    r"\b(?:the|a|of|for|if|is|are|find|so|with|when|"
    r"step|try|use|sum|verify|cross|check|or|and|to|from|"
    r"original|new|active|low|each|day|year|"
    r"AC|DC|via|cm|"
    r"primary|secondary)\b",
    re.IGNORECASE,
)


def _is_math_only_solve_line(text: str) -> bool:
    """True for indented solve lines that are pure math (no English prose)."""
    if not text.startswith("    "):
        return False
    if "=" not in text and "\\Rightarrow" not in text:
        return False
    body = text.strip()
    # Reject lines that contain prose words OUTSIDE of underscore subscripts.
    # First strip _{anything}, _word, ^{anything}.
    stripped = re.sub(r"_\{[^{}]*\}", "", body)
    stripped = re.sub(r"_[A-Za-z][A-Za-z0-9]*", "", stripped)
    stripped = re.sub(r"\^\{[^{}]*\}", "", stripped)
    if _PROSE_WORDS.search(stripped):
        return False
    return True


def _try_replace_known_equation(p) -> bool:
    text = p.text.strip()
    if text in EQUATION_LATEX:
        replace_with_math_zone(p, EQUATION_LATEX[text])
        return True
    return False


def _try_replace_solve_line(p) -> bool:
    text = p.text
    if not _is_math_only_solve_line(text):
        return False
    leading = re.match(r"^(\s*)", text).group(1)
    body = text.strip()
    try:
        latex = expression_to_latex(body)
        replace_with_math_zone(p, latex, prefix_text=leading)
        return True
    except Exception as exc:
        print(f"[warn] solve line failed: {body!r} -> {exc}")
        return False


def _apply_subscripts(p) -> bool:
    text = p.text
    if "_" not in text and "\u20D7" not in text:
        return False

    matches: list[tuple[int, int, str, str]] = []
    for m in UNDER_SUB.finditer(text):
        matches.append((m.start(), m.end(), m.group(1), m.group(2)))
    for m in VEC_SUB.finditer(text):
        base = m.group(1) + "\u20D7"
        matches.append((m.start(), m.end(), base, m.group(2)))
    if not matches:
        return False

    matches.sort(key=lambda x: (x[0], -x[1]))
    nonoverlap: list[tuple[int, int, str, str]] = []
    last_end = -1
    for start, end, base, sub in matches:
        if start >= last_end:
            nonoverlap.append((start, end, base, sub))
            last_end = end

    parts: list[tuple[str, bool]] = []
    pos = 0
    for start, end, base, sub in nonoverlap:
        if start > pos:
            parts.append((text[pos:start], False))
        parts.append((base, False))
        parts.append((sub, True))
        pos = end
    if pos < len(text):
        parts.append((text[pos:], False))

    remove_runs_and_math(p)
    for txt, is_sub in parts:
        append_text_run(p, txt, subscript=is_sub)
    return True


# =====================================================================================
# Driver
# =====================================================================================


def main() -> None:
    doc = Document(str(DOCX))

    n_eq = n_solve = n_sub = 0
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        if _try_replace_known_equation(p):
            n_eq += 1
            continue
        if _try_replace_solve_line(p):
            n_solve += 1
            continue
        if _apply_subscripts(p):
            n_sub += 1

    doc.save(str(DOCX))
    print(f"Equation paragraphs upgraded to math zones:    {n_eq}")
    print(f"Solve-step lines upgraded to math zones:        {n_solve}")
    print(f"Other paragraphs subscript-formatted:           {n_sub}")


if __name__ == "__main__":
    main()
