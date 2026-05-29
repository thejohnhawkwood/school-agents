"""Generate Biology homework from structured slide summaries.

Student-facing output:
- Short direction line once, then numbered questions (no meta-language, no OCR/slide-process wording).
- Teacher key + report retain slide traceability internally.

Input: slide-summaries.json produced by interpret_slides.py
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt


def set_default_font(doc: Document, size: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(size)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(16)


def clean(s: str) -> str:
    base = (s or "").replace("\ufffd", "'")
    base = base.replace("\u2019", "'").replace("\u2018", "'").replace("\u201c", '"').replace("\u201d", '"')
    return re.sub(r"\s+", " ", base).strip()


def filter_content_lines(slide: dict[str, Any]) -> list[str]:
    """Visible text suitable for composing questions (excludes titles, filler, credits)."""
    title = clean(slide.get("title", "")).lower()
    out: list[str] = []
    for x in slide.get("visible_text", []):
        line = clean(x)
        if not line:
            continue
        ll = line.lower()
        if title and ll == title:
            continue
        if ll in ("think-pair-share", "turn and talk"):
            continue
        if re.fullmatch(r"chapter\s+\d+(\.\d+)?", ll):
            continue
        if ll in ("unit d",):
            continue
        if "kid's health" in ll or "kids health" in ll:
            continue
        out.append(line)
    return out


def joined_blob(lines: list[str]) -> str:
    return " ".join(lines).lower()


def is_objectives_slide(slide: dict[str, Any]) -> bool:
    hay = " ".join(
        [
            clean(slide.get("title", "")),
            " ".join(clean(x) for x in slide.get("visible_text", [])),
        ]
    ).lower()
    return any(
        m in hay
        for m in (
            "learning outcomes",
            "learning and language objective",
            "swbat",
            "success criteria",
            "today you will",
        )
    )


def is_review_slide(slide: dict[str, Any]) -> bool:
    hay = " ".join(
        [
            clean(slide.get("title", "")),
            " ".join(clean(x) for x in slide.get("visible_text", [])),
        ]
    ).lower()
    return any(
        m in hay
        for m in (
            "review",
            "check your understanding",
            "exit ticket",
            "recap",
            "practice questions",
        )
    )


def has_instructional_visual(slide: dict[str, Any]) -> bool:
    return any(img.get("instructional_relevance") == "high" for img in slide.get("images", []))


def readable_alpha_ratio(text: str) -> float:
    letters = sum(ch.isalpha() for ch in text)
    total = len(text.replace(" ", ""))
    if total == 0:
        return 0.0
    return letters / total


def has_readable_instructional_ocr(slide: dict[str, Any]) -> bool:
    for img in slide.get("images", []):
        if img.get("instructional_relevance") != "high":
            continue
        ocr = clean(img.get("ocr_text", ""))
        if len(ocr) < 40:
            continue
        if readable_alpha_ratio(ocr) < 0.55:
            continue
        if len(ocr.split()) < 6:
            continue
        return True
    return False


def has_substantive_visible_text(slide: dict[str, Any]) -> bool:
    lines = filter_content_lines(slide)
    return len(" ".join(lines).split()) >= 8


def is_usable_evidence_slide(slide: dict[str, Any]) -> bool:
    if slide.get("confidence") == "low" and slide.get("flags"):
        return False
    return has_substantive_visible_text(slide) or has_readable_instructional_ocr(slide)


def short_join(parts: list[str], max_words: int) -> str:
    words: list[str] = []
    for p in parts:
        words.extend(clean(p).split())
        if len(words) >= max_words:
            break
    out = " ".join(words[:max_words])
    return out + (" ..." if len(words) > max_words else "")


def slide_to_student_question(slide: dict[str, Any]) -> str | None:
    """Return one student-facing question (1–2 sentences) or None if we should skip the slide."""
    if is_objectives_slide(slide):
        return None

    lines = filter_content_lines(slide)
    title_txt = clean(slide.get("title", ""))
    if not lines and not title_txt:
        return None

    # Include title in match text: filter_content_lines drops title duplicates, but many pattern cues live only there.
    blob = joined_blob([title_txt] + lines) if title_txt else joined_blob(lines)

    # --- Pattern-based questions (anchored to slide wording, not OCR) ---

    if "why does our body need an excretory system" in blob:
        return (
            "Why does the body need an excretory system? Relate your answer to cell metabolism and why "
            "some waste products can be harmful."
        )

    if (
        "excretion is the process of separating wastes from the body fluids and eliminating them" in blob
        or ("excretion" in blob and "separating wastes" in blob and "body fluids" in blob)
    ) and ("respiratory system and skin" in blob or "skin are both involved in excretion" in blob):
        return (
            "What is excretion? Explain how the respiratory system and skin can participate in excretion, "
            "and why eliminating feces is not considered excretion."
        )

    if "liver" in blob and "ammonia" in blob and "urea" in blob and "kidneys" in blob:
        return (
            "Explain how the liver processes nitrogen from proteins into urea, and what the kidneys do with "
            "urea and uric acid in the blood."
        )

    if (
        "ureters" in blob
        and "urethra" in blob
        and ("bladder" in blob or "urinary sphincter" in blob)
        and "kidneys" in blob
    ):
        return (
            "Trace the path of urine from where waste is filtered in the kidneys to where it leaves the body. "
            "Include the ureters, bladder, urinary sphincter, and urethra."
        )

    if "200 ml" in blob and "400 ml" in blob and "600 ml" in blob and "bladder" in blob:
        return (
            "How does the bladder signal the brain as urine volume increases, and what can happen to voluntary "
            "control at very high volume?"
        )

    if (
        "cortex" in blob
        and "medulla" in blob
        and "pelvis" in blob
        and "kidney" in blob
        and "ureter" in blob
    ):
        return (
            "Describe the cortex, medulla, and renal pelvis of the kidney. How does the renal pelvis connect "
            "the kidney to the ureter?"
        )

    if "functional units of the kidney" in blob and "glomerulus" in blob:
        return (
            "What is a nephron, and how does blood reach and leave the glomerulus through the afferent and "
            "efferent arterioles?"
        )

    if "bowman" in blob and "glomerulus" in blob and ("proximal tubule" in blob or "loop of henle" in blob):
        return (
            "Explain how filtrate moves from the glomerulus into Bowman's capsule and then through the proximal "
            "tubule, loop of Henle, and distal tubule."
        )

    # --- Urine formation in the nephron (typical Bio 20 deck) ---
    if "formation of urine involves three functions" in blob or (
        "filtration" in blob and "reabsorption" in blob and "secretion" in blob
    ):
        return (
            "Urine formation involves filtration, reabsorption, and secretion. In one or two sentences each, describe "
            "what each step accomplishes."
        )

    if "blood moves into the glomerulus" in blob and "bowman" in blob and "filtrate" in blob:
        return (
            "How does filtration at the glomerulus produce filtrate in Bowman's capsule, and why do plasma proteins "
            "and blood cells normally stay behind in the blood?"
        )

    if "120 ml" in blob and "reabsorbed" in blob and "dehydration" in blob:
        return (
            "About how much fluid is filtered each minute, and why is most of it reabsorbed? What risk does the body "
            "face if too much water stays in the filtrate?"
        )

    if "carrier molecules" in blob and "osmotic" in blob:
        return (
            "How can transport of ions out of the nephron change osmotic pressure and influence water movement?"
        )

    if "secretion is the movement of wastes from the blood into the nephron" in blob:
        return (
            "What is tubular secretion, and how can it move wastes from the blood into the filtrate?"
        )

    if "water balance" in blob and "endocrine" in blob:
        return (
            "When water intake changes, how can the kidneys adjust urine output, and why do they rely on nervous and "
            "endocrine signals?"
        )

    # --- Maintaining the system (ADH / aldosterone / pH) ---
    if "strictly regulating water" in blob and "adh" in blob and "aldosterone" in blob:
        return (
            "Why must the excretory system tightly regulate water, salts, and minerals, and how do ADH and aldosterone "
            "help maintain those balances?"
        )

    if "antidiuretic hormone" in blob or (
        "adh" in blob and "osmotic pressure" in blob and "water absorption" in blob
    ):
        return (
            "What does ADH change in the kidneys, and how does that affect whether urine becomes more concentrated?"
        )

    if "hypothalamus" in blob and "thirst" in blob:
        return (
            "How can blood solute levels affect hypothalamus cells, ADH release, and the sensation of thirst?"
        )

    if "as you take in water" in blob and "adh" in blob:
        return (
            "When you drink water, what happens to blood solute levels and ADH release, and how does that change water "
            "reabsorption in the nephrons?"
        )

    if "adh and the nephron" in clean(slide.get("title", "")).lower() or (
        "proximal tubule" in blob and "collecting duct" in blob and "adh" in blob
    ):
        return (
            "Where is most filtered water reabsorbed along the nephron, and how does ADH change water movement in the "
            "distal tubule and collecting duct?"
        )

    if "aldosterone" in blob and "blood pressure" in blob:
        return (
            "How does aldosterone affect sodium reabsorption in the nephrons, and why can that influence blood volume "
            "and blood pressure?"
        )

    if "acid-base buffer" in blob or ("blood ph" in blob and "kidneys" in blob):
        return (
            "Why does blood pH need buffering, and how do hydrogen and bicarbonate ions factor into acid-base balance?"
        )

    if "kidneys help to maintain this balance" in blob and "hco" in blob.replace(" ", ""):
        return (
            "How can the kidneys help balance hydrogen and bicarbonate levels when blood becomes too acidic?"
        )

    # --- Disorders / dialysis ---
    if "diabetes mellitus" in blob and "insulin" in blob:
        return (
            "How does inadequate insulin raise blood sugar, and why can that lead to extra urine volume and sugar in "
            "the urine?"
        )

    if "diabetes insipidus" in blob and "adh" in blob:
        return (
            "What goes wrong with ADH in diabetes insipidus, and why can daily urine output become extremely high?"
        )

    if "nephritis" in blob.lower() or "bright" in blob.lower():
        return (
            "How can inflammation of nephrons let proteins enter the filtrate, and why might urine production increase?"
        )

    if "hemodialysis" in blob and "semipermeable" in blob:
        return (
            "How does hemodialysis use a semipermeable membrane and dialysis solution to remove wastes from blood?"
        )

    if "peritoneal dialysis" in blob and "peritoneum" in blob:
        return (
            "How does peritoneal dialysis use the peritoneum as a dialysis surface, and where is the dialysate placed?"
        )

    if "dialysis" in blob and "machine" in blob and "filtered" in blob:
        return (
            "When kidneys fail, why can dialysis be life-saving, and what is the basic idea of filtering blood through "
            "a machine?"
        )

    if "analyze urine" in blob and "diabetes" in blob:
        return (
            "Why might healthcare providers consider diet, stress, and activity when interpreting urine tests for "
            "signs of disease?"
        )

    if "dialysis treatment takes between 2 and 5 hours" in blob or (
        "hemodialysis" in blob and "strictly regulated diet" in blob
    ):
        return (
            "About how long does a hemodialysis treatment usually last, how often might it be repeated each week, and "
            "why must diet be strictly regulated?"
        )

    if "kidney transplants" in blob and "family member donating" in blob:
        return (
            "Why is a kidney transplant described as a permanent alternative to dialysis, and why might a close "
            "family member donate a kidney?"
        )

    if "kidney transplants" in blob and "cadavers" in blob:
        return (
            "Why should a donor and recipient monitor their diet after donating a kidney, and how can kidneys from "
            "cadavers be used?"
        )

    if "kidney-coronary connection" in clean(slide.get("title", "")).lower() or (
        "high blood pressure" in blob and "blood vessels in the kidneys" in blob
    ):
        return (
            "How can high blood pressure damage kidney blood vessels and reduce waste filtering, and why might "
            "symptoms show up only after serious damage?"
        )

    if "kidney stones" in blob and "mineral salts" in blob:
        return (
            "What are kidney stones made of, where can they lodge and cause pain, and how can ultrasound help break "
            "them up?"
        )

    # Explicit slide questions (keep concise)
    for line in lines:
        if line.endswith("?"):
            return line

    # Last resort: strongest complete-looking sentence from visible text (avoid leading fragments such as "As well,")
    bad_prefixes = ("as well", "however,", "therefore,", "these ", "this ")
    substantive = []
    for ln in lines:
        low = ln.strip().lower()
        if len(ln.split()) < 8 or low.startswith("swbat"):
            continue
        if any(low.startswith(p) for p in bad_prefixes):
            continue
        substantive.append(ln)
    if not substantive:
        substantive = [
            ln for ln in lines if len(ln.split()) >= 8 and not ln.lower().startswith("swbat")
        ]
    if substantive:
        core = sorted(substantive, key=lambda x: len(x.split()), reverse=True)[0].rstrip(".")
        return f"Explain {core}."

    return None


def score_slide_for_selection(slide: dict[str, Any]) -> tuple[int, list[str]]:
    reasons: list[str] = []
    score = 0
    conf = slide.get("confidence", "medium")
    if conf == "low":
        score += 6
        reasons.append("low-confidence slide")
    elif conf == "medium" and not has_substantive_visible_text(slide):
        score += 2
        reasons.append("medium-confidence slide with thin visible text")

    if not is_usable_evidence_slide(slide):
        score += 15
        reasons.append("thin evidence")

    if is_review_slide(slide) and not has_instructional_visual(slide):
        score += 20
        reasons.append("review slide")

    if slide_to_student_question(slide) is None:
        score += 25
        reasons.append("no clean student question derived")

    return score, reasons


def pick_standard_slides(slides: list[dict[str, Any]], count: int = 5) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    scored = []
    for idx, slide in enumerate(slides):
        s, reasons = score_slide_for_selection(slide)
        scored.append((s, idx, slide, reasons))
    scored.sort(key=lambda x: (x[0], x[1]))

    picked: list[dict[str, Any]] = []
    notes: list[dict[str, Any]] = []
    for s, _, slide, reasons in scored:
        if len(picked) >= count:
            break
        if is_objectives_slide(slide):
            continue
        if is_review_slide(slide) and not has_instructional_visual(slide):
            continue
        if not is_usable_evidence_slide(slide):
            continue
        q = slide_to_student_question(slide)
        if q is None:
            continue
        picked.append(slide)
        notes.append({"source_ref": slide.get("source_ref", "?"), "score": s, "notes": reasons})
    return picked, notes


def pick_diagram_slide(slides: list[dict[str, Any]]) -> dict[str, Any] | None:
    """Prefer urinary-system or nephron slides with solid text (verified content for the prompt)."""
    if not slides:
        return None

    def title_of(s: dict[str, Any]) -> str:
        return clean(s.get("title", "")).lower()

    for s in slides:
        if "urinary system" in title_of(s):
            return s
    for s in slides:
        lines = filter_content_lines(s)
        b = joined_blob(lines)
        if "ureters" in b and "urethra" in b and "kidneys" in b:
            return s
    for s in slides:
        if "nephron" in title_of(s) and has_substantive_visible_text(s):
            return s
    for s in slides:
        t = title_of(s)
        if ("hemodialysis" in t or "dialysis" == t or t.startswith("dialysis")) and has_substantive_visible_text(s):
            return s
    for s in slides:
        if has_instructional_visual(s) and has_substantive_visible_text(s):
            return s
    return slides[0]


def student_diagram_prompt(diagram_slide: dict[str, Any]) -> str:
    lines = filter_content_lines(diagram_slide)
    b = joined_blob(lines)
    t = clean(diagram_slide.get("title", "")).lower()

    if "urinary system" in t or ("ureters" in b and "urethra" in b):
        return (
            "Draw the urinary system. Label the kidneys, ureters, urinary bladder, and urethra. Beside your "
            "diagram, summarize how urine moves from the kidneys out of the body."
        )

    if "nephron" in t or ("glomerulus" in b and "bowman" in b):
        return (
            "Draw and label a nephron. Show the glomerulus, Bowman's capsule, and the tubule route through the "
            "proximal tubule, loop of Henle, and distal tubule."
        )

    if "kidney" in b and "cortex" in b and "medulla" in b:
        return (
            "Sketch a kidney showing the cortex, medulla, and renal pelvis. Label each region and show where "
            "the ureter attaches."
        )

    if "dialysis" in t and ("machine" in b or "blood is filtered" in b):
        return (
            "Draw a simple dialysis setup showing blood leaving the body, passing through a machine that filters "
            "wastes, and returning to the body. Label the machine and the direction of blood flow."
        )

    if "filtration" in b and "reabsorption" in b and "secretion" in b:
        return (
            "Make a simple three-panel sketch for filtration, reabsorption, and secretion. Label each panel and "
            "write one sentence describing what moves between blood, filtrate, and the nephron."
        )

    if "semipermeable" in b or "hemodialysis" in t:
        return (
            "Sketch hemodialysis as blood passing through tubing next to dialysis fluid, separated by a semipermeable "
            "membrane. Label where wastes diffuse into the surrounding solution."
        )

    if "peritoneal dialysis" in t:
        return (
            "Draw the abdomen showing dialysate in the peritoneal cavity and indicate how waste exchange occurs across "
            "the peritoneum."
        )

    return (
        "Draw and label the main structures taught on this lesson's key diagram, and write two sentences "
        "explaining the process it illustrates."
    )


def teacher_answer_stub(slide: dict[str, Any], student_q: str) -> str:
    src = slide.get("source_ref", "?")
    lines = filter_content_lines(slide)
    return (
        f"Source {src}. Expect answers grounded in: {short_join(lines, 45)}. "
        f"Student prompt: {short_join([student_q], 40)}"
    )


def write_student_docx(
    path: Path,
    lesson_title: str,
    questions: list[str],
    diagram_prompt: str,
) -> None:
    doc = Document()
    set_default_font(doc)
    add_title(doc, lesson_title)
    doc.add_paragraph("Answer in complete sentences.")
    doc.add_paragraph("")
    for i, q in enumerate(questions, start=1):
        doc.add_paragraph(f"{i}. {q}")
        doc.add_paragraph("")
    pb = doc.add_paragraph()
    pb.add_run().add_break(WD_BREAK.PAGE)
    head = doc.add_paragraph()
    hr = head.add_run("6. Diagram question")
    hr.bold = True
    doc.add_paragraph(diagram_prompt)
    doc.save(path)


def write_teacher_key_docx(path: Path, title: str, bullets: list[str]) -> None:
    doc = Document()
    set_default_font(doc)
    add_title(doc, title)
    for i, text in enumerate(bullets, start=1):
        p = doc.add_paragraph()
        r = p.add_run(f"{i}. ")
        r.bold = True
        p.add_run(text)
    doc.save(path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate homework from slide summaries.")
    parser.add_argument("--summary-json", required=True)
    parser.add_argument("--lesson-title", required=True)
    parser.add_argument("--outdir", required=True)
    parser.add_argument("--run-id", required=True)
    args = parser.parse_args()

    payload = json.loads(Path(args.summary_json).read_text(encoding="utf-8"))
    slides: list[dict[str, Any]] = payload.get("slides", [])
    allowed_range = payload.get("allowed_range", "unknown")
    source_file = payload.get("source_file", "unknown")

    picked, selection_notes = pick_standard_slides(slides, count=5)
    issues: list[dict[str, Any]] = []

    if len(picked) < 5:
        issues.append(
            {
                "question": "all",
                "source_ref": allowed_range,
                "reason": "fewer than 5 slides yielded clean student questions under current rules",
                "action": "narrow range or improve slide text extraction",
            }
        )

    student_questions: list[str] = []
    teacher_lines: list[str] = []
    mapping: dict[str, list[str]] = {}

    for i, slide in enumerate(picked[:5], start=1):
        q = slide_to_student_question(slide)
        assert q is not None
        student_questions.append(q)
        src = slide.get("source_ref", "?")
        mapping[str(i)] = [src]
        teacher_lines.append(teacher_answer_stub(slide, q))

    diagram_slide = pick_diagram_slide(slides)
    diagram_q = student_diagram_prompt(diagram_slide) if diagram_slide else "Diagram question unavailable."
    if diagram_slide:
        mapping["6"] = [diagram_slide.get("source_ref", "?")]
        teacher_lines.append(
            f"Diagram graded for completeness and labeling consistent with {diagram_slide.get('source_ref')}. "
            "Expect kidneys → ureters → bladder → urethra for urinary-system prompts, or nephron regions for nephron prompts."
        )
    else:
        mapping["6"] = []
        teacher_lines.append("No diagram slide selected.")

    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)

    student_doc = outdir / f"{args.run_id}-student.docx"
    teacher_doc = outdir / f"{args.run_id}-teacher-key.docx"
    report_md = outdir / f"{args.run_id}-report.md"
    audit_json = outdir / f"{args.run_id}-audit.json"
    issues_json = outdir / f"{args.run_id}-issues.json"

    write_student_docx(student_doc, args.lesson_title, student_questions, diagram_q)
    write_teacher_key_docx(teacher_doc, f"{args.lesson_title} - Teacher Key", teacher_lines)

    report_lines = [
        f"# {args.lesson_title} - Generation Report",
        "",
        f"- Source file: `{Path(source_file).name}`",
        f"- Allowed range: `{allowed_range}`",
        f"- Questions generated: {len(student_questions) + 1}",
        f"- Flags: {len(issues)}",
        "",
        "## Question-to-source mapping",
    ]
    for qid, refs in mapping.items():
        report_lines.append(f"- {qid} -> {', '.join(refs)}")
    report_lines.extend(["", "## Source selection QA"])
    for note in selection_notes:
        note_text = ", ".join(note["notes"]) if note["notes"] else "preferred source"
        report_lines.append(f"- {note['source_ref']}: score={note['score']} ({note_text})")
    report_lines.extend(["", "## Review flags"])
    if not issues:
        report_lines.append("- None")
    else:
        for item in issues:
            report_lines.append(
                f"- Q{item.get('question')}: {item['reason']} ({item['source_ref']}) -> {item['action']}"
            )

    report_md.write_text("\n".join(report_lines), encoding="utf-8")

    audit_payload = {
        "source_file": source_file,
        "allowed_range": allowed_range,
        "question_to_source": mapping,
        "selection_notes": selection_notes,
        "slide_count": len(slides),
        "student_questions_preview": student_questions + [diagram_q],
        "qa_decision": "pass-with-flags" if issues else "pass",
    }
    audit_json.write_text(json.dumps(audit_payload, indent=2), encoding="utf-8")
    issues_json.write_text(json.dumps({"issues": issues}, indent=2), encoding="utf-8")

    print(f"Wrote: {student_doc}")
    print(f"Wrote: {teacher_doc}")
    print(f"Wrote: {report_md}")
    print(f"Wrote: {audit_json}")
    print(f"Wrote: {issues_json}")


if __name__ == "__main__":
    main()
