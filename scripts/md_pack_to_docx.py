"""Convert Mars research pack Markdown (student packets) to Word .docx.

Handles: # to ### headings, paragraphs, ordered lists (1. 2. …), bullet lists,
checkbox bullets, markdown tables, fenced code blocks, horizontal rules,
blockquotes, **bold**, `code`, and [text](url) links.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_inline_runs(paragraph, text: str) -> None:
    """Split **bold**, `code`, and [label](url) into runs."""
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        elif part.startswith("[") and "](" in part:
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if m:
                label, url = m.group(1), m.group(2)
                paragraph.add_run(f"{label} ({url})")
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)


def is_table_separator(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|"):
        return False
    return bool(re.match(r"^\|[\s\-:|]+\|$", s))


def flush_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell_text = row[ci] if ci < len(row) else ""
            tbl.rows[ri].cells[ci].text = cell_text


def convert(md_path: Path, out_path: Path) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_fence = False
    fence_buf: list[str] = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip("\n")

        if line.strip().startswith("```"):
            if not in_fence:
                in_fence = True
                fence_buf = []
            else:
                in_fence = False
                p = doc.add_paragraph()
                run = p.add_run("\n".join(fence_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                fence_buf = []
            i += 1
            continue

        if in_fence:
            fence_buf.append(raw)
            i += 1
            continue

        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        if "|" in line and line.strip().startswith("|"):
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_line = lines[i]
                if is_table_separator(row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.split("|")[1:-1]]
                rows.append(cells)
                i += 1
            flush_table(doc, rows)
            continue

        # Markdown ordered lists: "1. ", "2. ", … (must not match headings — those use #)
        num_m = re.match(r"^\s*(\d+)\.\s+(.+)$", line)
        if num_m:
            while i < len(lines):
                ln = lines[i]
                m2 = re.match(r"^\s*(\d+)\.\s+(.+)$", ln)
                if not m2:
                    break
                item_text = m2.group(2).strip()
                p = doc.add_paragraph(style="List Number")
                add_inline_runs(p, item_text)
                i += 1
            continue

        if line.strip().startswith("- [ ]"):
            text = line.strip()[5:].strip()
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("☐ ")
            add_inline_runs(p, text)
            i += 1
            continue

        if line.strip().startswith("- "):
            text = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, text)
            i += 1
            continue

        if line.startswith(">"):
            p = doc.add_paragraph()
            add_inline_runs(p, line[1:].strip())
            p.paragraph_format.left_indent = Pt(24)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        # Merge continuation lines into one paragraph until blank or special
        para_parts = [line.strip()]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            if not nxt.strip():
                break
            if nxt.startswith("#"):
                break
            if nxt.strip().startswith("```"):
                break
            if nxt.strip().startswith("|"):
                break
            if nxt.strip().startswith("- "):
                break
            if re.match(r"^\s*\d+\.\s+", nxt.strip()):
                break
            if nxt.startswith(">"):
                break
            para_parts.append(nxt.strip())
            i += 1
        p = doc.add_paragraph()
        add_inline_runs(p, " ".join(para_parts))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("md_path", type=Path)
    parser.add_argument("-o", "--out", type=Path, default=None)
    args = parser.parse_args()
    out = args.out or args.md_path.with_suffix(".docx")
    convert(args.md_path, out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
