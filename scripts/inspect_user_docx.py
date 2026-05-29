"""Inspect the user-formatted DOCX so we can plan surgical math upgrades
without disturbing column / page-break formatting."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
from pathlib import Path
from docx import Document

SRC = Path(r"school-agents/outputs/homework/sci30-unitc-emr-review-v1/sci30-unitc-emr-review-student.docx")

doc = Document(str(SRC))

print(f"Sections: {len(doc.sections)}")
for i, sec in enumerate(doc.sections):
    sectPr = sec._sectPr
    cols = sectPr.find(f"{{{sectPr.nsmap['w']}}}cols") if 'w' in sectPr.nsmap else None
    print(f"  Section {i}: page_w={sec.page_width}, page_h={sec.page_height}, "
          f"cols_attrs={None if cols is None else dict(cols.attrib)}")

print(f"\nParagraphs: {len(doc.paragraphs)}")
print("\nFirst 80 paragraphs (style, has_pb, text preview):")
for i, p in enumerate(doc.paragraphs[:80]):
    style = p.style.name
    text = p.text[:100].replace("\n", " | ")
    has_pb = any("pageBreakBefore" in str(r._r.xml) or "<w:br" in str(r._r.xml) for r in p.runs)
    print(f"  {i:3d}  [{style:15s}] pb={int(has_pb)}  {text}")

print(f"\n... total {len(doc.paragraphs)} paragraphs ...")
print("\nTail (last 10):")
for i, p in enumerate(doc.paragraphs[-10:]):
    print(f"  -{10-i:2d}  [{p.style.name:15s}]  {p.text[:120]}")
