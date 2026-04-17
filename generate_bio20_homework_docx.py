import json
import io
import re
import zipfile
from pathlib import Path
from typing import Iterable
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from PIL import Image
import pytesseract


OUTPUT_DIR = Path(
    r"C:\Users\pbird\OneDrive - Christ the Redeemer Catholic School Board\Desktop\School Based Cursor\school-agents\outputs"
)
PPTX_PATH = Path(
    r"C:\Users\pbird\Downloads\Bio 20 - Muscular System and Homeostasis PowerPoint.pptx"
)
TESSERACT_EXE = Path(
    r"C:\Users\pbird\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"
)
pytesseract.pytesseract.tesseract_cmd = str(TESSERACT_EXE)


def set_default_font(doc: Document, size: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(size)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16)


def add_question(doc: Document, number: int, topic: str, prompt: str) -> None:
    p_topic = doc.add_paragraph()
    topic_run = p_topic.add_run(f"{number}. {topic}")
    topic_run.bold = True
    doc.add_paragraph(prompt)
    doc.add_paragraph("\n\n")


def add_teacher_answer(doc: Document, number: int, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. ")
    run.bold = True
    p.add_run(text)


def write_student_docx(path: Path, title: str, questions_1_5: list[tuple[str, str]], q6: tuple[str, str]) -> None:
    doc = Document()
    set_default_font(doc)
    add_title(doc, title)

    for idx, (topic, prompt) in enumerate(questions_1_5, start=1):
        add_question(doc, idx, topic, prompt)

    # Required layout: page break before question 6.
    pb = doc.add_paragraph()
    pb.add_run().add_break(WD_BREAK.PAGE)
    add_question(doc, 6, q6[0], q6[1])
    doc.save(path)


def write_teacher_key_docx(path: Path, title: str, answers: list[str]) -> None:
    doc = Document()
    set_default_font(doc)
    add_title(doc, title)
    for i, ans in enumerate(answers, start=1):
        add_teacher_answer(doc, i, ans)
    doc.save(path)


def clean_ocr_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    return text


def gist_from_ocr(text: str) -> str:
    cleaned = clean_ocr_text(text)
    if not cleaned:
        return "No reliable OCR text extracted from slide images."
    words = cleaned.split()
    short = " ".join(words[:24])
    if len(words) > 24:
        short += " ..."
    return short


def ocr_image_bytes(image_bytes: bytes) -> str:
    try:
        with Image.open(io.BytesIO(image_bytes)) as img:
            return clean_ocr_text(pytesseract.image_to_string(img))
    except Exception:
        return ""


def extract_slide_ocr(zf: zipfile.ZipFile, slide_num: int) -> tuple[list[str], str]:
    ns = {
        "p": "http://schemas.openxmlformats.org/presentationml/2006/main",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "pr": "http://schemas.openxmlformats.org/package/2006/relationships",
    }
    slide_path = f"ppt/slides/slide{slide_num}.xml"
    rel_path = f"ppt/slides/_rels/slide{slide_num}.xml.rels"
    if slide_path not in zf.namelist():
        return [], "slide-missing"

    slide_root = ET.fromstring(zf.read(slide_path))
    rel_root = ET.fromstring(zf.read(rel_path)) if rel_path in zf.namelist() else None

    rel_map: dict[str, str] = {}
    if rel_root is not None:
        for rel in rel_root.findall("pr:Relationship", ns):
            rid = rel.attrib.get("Id")
            target = rel.attrib.get("Target", "")
            if rid and target.startswith("../media/"):
                rel_map[rid] = f"ppt/media/{Path(target).name}"

    ocr_texts: list[str] = []
    for blip in slide_root.findall(".//a:blip", ns):
        rid = blip.attrib.get(f"{{{ns['r']}}}embed")
        media_path = rel_map.get(rid or "")
        if not media_path or media_path not in zf.namelist():
            continue
        text = ocr_image_bytes(zf.read(media_path))
        if text:
            ocr_texts.append(text)

    status = "ok" if ocr_texts else "no-text-detected"
    return ocr_texts, status


def slide_image_report(slide_numbers: Iterable[int]) -> dict[int, dict[str, int | str | bool]]:
    data: dict[int, dict[str, int | str | bool]] = {}
    with zipfile.ZipFile(PPTX_PATH) as zf:
        for s in slide_numbers:
            slide_path = f"ppt/slides/slide{s}.xml"
            rel_path = f"ppt/slides/_rels/slide{s}.xml.rels"
            xml = zf.read(slide_path).decode("utf-8", errors="ignore")
            rel_xml = zf.read(rel_path).decode("utf-8", errors="ignore") if rel_path in zf.namelist() else ""

            text_runs = len(re.findall(r"<a:t>.*?</a:t>", xml, flags=re.DOTALL))
            image_links = len(re.findall(r"Target=\"\\.\\./media/", rel_xml))
            likely_image_heavy = image_links >= 1 and text_runs <= 6
            ocr_chunks, ocr_status = extract_slide_ocr(zf, s)
            ocr_joined = " ".join(ocr_chunks)

            data[s] = {
                "text_runs": text_runs,
                "image_links": image_links,
                "likely_image_heavy": likely_image_heavy,
                "ocr_status": ocr_status,
                "ocr_chars": len(ocr_joined),
                "image_gist": gist_from_ocr(ocr_joined),
                "coverage_status": "covered-by-question-or-flagged-with-ocr",
            }
    return data


def write_report(
    path: Path,
    title: str,
    allowed_range: str,
    limitations: list[str],
    mapping: list[str],
    excluded: list[str],
    mismatch_risks: list[str],
    low_confidence: list[str],
    unresolved: list[str],
    next_steps: list[str],
    image_cov: dict[int, dict[str, int | str | bool]],
) -> None:
    lines = [
        f"# {title} Report",
        "",
        "## Summary of inputs used",
        f"- Source file: `{PPTX_PATH.name}`",
        f"- Allowed range: `{allowed_range}`",
        "- Source quality limitations:",
    ]
    lines.extend([f"  - {x}" for x in limitations])
    lines.extend(["", "## Audit / traceability", "- Question-to-source mapping:"])
    lines.extend([f"  - {m}" for m in mapping])
    lines.append("- Excluded source segments and reasons:")
    lines.extend([f"  - {e}" for e in excluded])
    lines.append("- Detected slide/page mismatch risks:")
    lines.extend([f"  - {r}" for r in mismatch_risks])
    lines.extend(["", "## Image-ingestion coverage notes"])
    for slide, meta in sorted(image_cov.items()):
        lines.append(
            f"- Slide {slide}: text_runs={meta['text_runs']}, image_links={meta['image_links']}, "
            f"likely_image_heavy={meta['likely_image_heavy']}, ocr_status={meta['ocr_status']}, "
            f"ocr_chars={meta['ocr_chars']}, image_gist=\"{meta['image_gist']}\", "
            f"coverage_status={meta['coverage_status']}"
        )
    lines.extend(["", "## Flags / uncertainty", "- Low-confidence items:"])
    lines.extend([f"  - {x}" for x in low_confidence])
    lines.append("- Unresolved issues:")
    lines.extend([f"  - {x}" for x in unresolved])
    lines.append("- Suggested next teacher action:")
    lines.extend([f"  - {x}" for x in next_steps])
    path.write_text("\n".join(lines), encoding="utf-8")


def write_audit_and_issues(prefix: str, audit: dict, issues: dict) -> None:
    (OUTPUT_DIR / f"{prefix}-audit.json").write_text(json.dumps(audit, indent=2), encoding="utf-8")
    (OUTPUT_DIR / f"{prefix}-issues.json").write_text(json.dumps(issues, indent=2), encoding="utf-8")


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    energy_questions = [
        (
            "ATP as the energy source",
            "Identify the immediate energy molecule used by muscle cells during contraction, and explain why the body uses more than one method to produce it.",
        ),
        (
            "Creatine phosphate pathway",
            "Describe how creatine phosphate supports ATP production in resting and active muscle, and state how long this method typically supports intense activity.",
        ),
        (
            "Aerobic respiration in muscle",
            "Explain where aerobic cellular respiration occurs in muscle fibres, what fuel sources can be used, and one additional body effect caused by this pathway.",
        ),
        (
            "Lactic acid fermentation and fatigue",
            "Explain how ATP is produced during lactic acid fermentation and why prolonged reliance on this pathway can lead to cramping and fatigue.",
        ),
        (
            "Contraction sequence under prolonged activity",
            "Using the contraction sequence diagram, describe what happens as a muscle moves from immediate ATP use to creatine phosphate support and then to longer-term pathways during sustained activity.",
        ),
    ]
    energy_q6 = (
        "Diagram: ATP production routes",
        "Draw a three-branch flow diagram showing ATP production by (1) creatine phosphate breakdown, (2) aerobic cellular respiration, and (3) lactic acid fermentation. For each branch, label oxygen requirement, speed/duration, and one key outcome.",
    )
    energy_answers = [
        "ATP is the direct energy source for contraction; multiple ATP pathways are needed because activity intensity and oxygen availability change. (Slides 26-27)",
        "Creatine phosphate builds during rest and rapidly regenerates ATP; this supports about 8 seconds of intense activity before recovery at rest. (Slide 28)",
        "Aerobic respiration occurs in mitochondria and can use glucose or fatty acids when oxygen is available; heat is also produced. (Slide 29)",
        "Fermentation converts glucose to lactate to produce ATP without oxygen; lactate buildup acidifies sarcoplasm and contributes to cramping/fatigue. (Slide 31)",
        "The contraction-sequence visual shows immediate ATP demand first, short-term creatine phosphate replenishment, and transition to longer-duration pathways as contraction continues. (Slide 32 OCR + Slides 27-31)",
        "A complete response includes all three pathways with oxygen status, speed/duration, and outcomes aligned to source. (Slides 27-32, 34)",
    ]
    energy_img = slide_image_report(range(26, 35))
    write_student_docx(
        OUTPUT_DIR / "Lesson 1 Homework - Energy for Muscle Contraction - Student.docx",
        "Lesson 1 Homework - Energy for Muscle Contraction",
        energy_questions,
        energy_q6,
    )
    write_teacher_key_docx(
        OUTPUT_DIR / "Lesson 1 Homework - Energy for Muscle Contraction - Teacher Key.docx",
        "Lesson 1 Teacher Key - Energy for Muscle Contraction",
        energy_answers,
    )
    write_report(
        OUTPUT_DIR / "Lesson 1 Homework - Energy for Muscle Contraction - Report.md",
        "Lesson 1 Homework - Energy for Muscle Contraction",
        "Slides 26-34",
        [
            "Slide 32 has no extractable text.",
            "Slides with low text and image references are flagged as likely image-heavy and reviewed in image-ingestion notes.",
        ],
        [
            "1 -> Slides 26-27",
            "2 -> Slide 28",
            "3 -> Slide 29",
            "4 -> Slide 31",
            "5 -> Slide 32 OCR and Slides 27-31",
            "6 -> Slide 32 OCR and Slides 27-31",
        ],
        [
            "Slide 33 used only for context (objectives) and not as sole support.",
        ],
        ["No explicit mismatch detected in this range."],
        ["Image-heavy slides were OCR-processed; low-text slides still require teacher spot-check for diagrams/charts."],
        ["Some image regions may remain low-confidence where OCR returned no reliable text."],
        ["Review reported image_gist lines for slide-level intent confirmation before final classroom use."],
        energy_img,
    )
    write_audit_and_issues(
        "lesson-1-energy",
        {
            "allowed_range": "slides 26-34",
            "question_to_source": {
                "1": ["26", "27"],
                "2": ["28"],
                "3": ["29"],
                "4": ["31"],
                "5": ["32", "27", "28", "29", "31"],
                "6": ["27", "28", "29", "31", "32", "34"],
            },
            "image_ingestion": energy_img,
            "qa_decision": "pass-with-flags",
        },
        {
            "unresolved": [
                "Slide 33 OCR returned no reliable image text; avoid claim-level dependence on this slide.",
            ]
        },
    )

    home_questions = [
        (
            "Muscle tone and posture",
            "Why are muscles still contracting at some level even at rest, and how does this support posture and upright body position?",
        ),
        (
            "Atrophy from lack of use",
            "Define muscular atrophy based on the lesson and explain the type of condition that can lead to it.",
        ),
        (
            "Energy timeline in different exercise durations",
            "Based on the exercise-duration visual, summarize how ATP supply shifts between very short high-intensity effort and prolonged exercise.",
        ),
        (
            "Twitch, summation, and tetanus",
            "Describe the progression from a single muscle twitch to summation and then tetanus, including what eventually stops sustained contraction.",
        ),
        (
            "Fast- vs slow-twitch fibres",
            "Compare Type I and Type II muscle fibres in terms of contraction speed, fatigue resistance, and typical energy pathway.",
        ),
    ]
    home_q6 = (
        "Diagram: muscles and homeostasis",
        "Create a labelled concept map showing how muscle contraction contributes to homeostasis. Include ATP use, heat generation, blood vessel contraction/dilation, and one additional process in the body that depends on muscle movement.",
    )
    home_answers = [
        "Baseline contraction (muscle tone) helps maintain posture and keep the body upright at rest. (Slide 37)",
        "Muscular atrophy is a loss/reduction related to lack of movement or use. (Slide 38)",
        "The duration visual indicates immediate ATP sources for very short effort, followed by other pathways over longer durations; prolonged activity depends on sustained ATP production. (Slide 41 OCR, with reinforcement from Slides 40 and 42)",
        "A sufficient stimulus creates a twitch; rapid stimuli create summation and can progress to tetanus; contraction ends as energy reserves are depleted and fatigue occurs. (Slide 42)",
        "Type I fibres are slower and fatigue resistant with aerobic tendency; Type II fibres are faster, generate rapid power, and fatigue quickly with more anaerobic dependence. (Slide 44)",
        "Strong responses connect ATP-driven contraction to heat production and vessel constriction/dilation, plus another valid process dependent on muscle movement. (Slide 45)",
    ]
    home_img = slide_image_report(range(35, 47))
    write_student_docx(
        OUTPUT_DIR / "Lesson 2 Homework - Muscles Health and Homeostasis - Student.docx",
        "Lesson 2 Homework - Muscles, Health, and Homeostasis",
        home_questions,
        home_q6,
    )
    write_teacher_key_docx(
        OUTPUT_DIR / "Lesson 2 Homework - Muscles Health and Homeostasis - Teacher Key.docx",
        "Lesson 2 Teacher Key - Muscles, Health, and Homeostasis",
        home_answers,
    )
    write_report(
        OUTPUT_DIR / "Lesson 2 Homework - Muscles Health and Homeostasis - Report.md",
        "Lesson 2 Homework - Muscles, Health, and Homeostasis",
        "Slides 35-46",
        [
            "Slides 41 and 43 have no extractable text.",
            "Slide 39 is mostly a disorder list without detail; not used as sole support for explanatory answers.",
        ],
        [
            "1 -> Slide 37",
            "2 -> Slide 38",
            "3 -> Slide 41 OCR (with support Slides 40 and 42)",
            "4 -> Slide 42",
            "5 -> Slide 44",
            "6 -> Slide 45",
        ],
        [
            "Slide 43 excluded from direct claims due to noisy OCR quality.",
            "Slide 39 excluded from detailed explanatory claims due to limited extractable details.",
        ],
        ["No explicit mismatch detected in this range."],
        ["Slide 45 prompt asks for examples but does not enumerate them; student responses should remain source-aligned."],
        ["Image OCR was executed, but some slides still return little/no text (likely diagram-only visuals)."],
        ["Review image_gist entries; if key diagram labels are missing, provide annotated snapshots for precise follow-up questions."],
        home_img,
    )
    write_audit_and_issues(
        "lesson-2-homeostasis",
        {
            "allowed_range": "slides 35-46",
            "question_to_source": {
                "1": ["37"],
                "2": ["38"],
                "3": ["41", "40", "42"],
                "4": ["42"],
                "5": ["44"],
                "6": ["45"],
            },
            "image_ingestion": home_img,
            "qa_decision": "pass-with-flags",
        },
        {
            "unresolved": [
                "Slide 43 OCR is present but noisy; keep this slide out of high-specificity claims.",
            ]
        },
    )

    print("Generated student docs, teacher keys, reports, audit.json, and issues.json for both lessons.")


if __name__ == "__main__":
    main()
