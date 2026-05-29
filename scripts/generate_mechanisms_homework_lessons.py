"""Generate Bio 20 homework DOCX bundles for Mechanisms deck lessons 5-8.

Student-facing stems avoid slide numbers and notepack-layout cues; topics anchor
questions to lesson content. Slide ranges stay in run folder names + audit JSON
paths for teacher/repo traceability only.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PARENT = ROOT / "outputs" / "homework"
DOC_VERSION = "v2"

LESSONS = [
    {
        "run_id": "bio20-mechanisms-lesson5-slides8-27-v2",
        "lesson": 5,
        "slide_span": "8-27",
        "topic": "Adaptations, Variation, Mutations and Natural Selection",
        "processed": "bio20-mechanisms-lesson5-slides8-27",
        "questions": [
            "According to the lesson, why are reproduction and survival tied to passing traits to the next generation, and why do the best-adapted organisms tend to survive? What three kinds of adaptations does the lesson contrast—structural, behavioural, and physiological—and how does the lesson define each type?",
            "The lesson sorts practice examples of adaptations into structural, physiological, and behavioural categories (including thick fur on bears; hair on mammal skin reducing sun exposure; blood-vessel changes or shivering for temperature balance; behaviours like wrapping up when cold or hibernating). Choose two structural examples, two physiological examples, and two behavioural examples from those lesson lists and explain briefly why each one fits its category.",
            "What does the lesson say causes adaptations to accumulate in a population over time (the gradual-change idea)? Where does inherited variation ultimately come from, and according to the lesson what can introduce DNA mutations (such as copying errors plus damage from radiation or chemical mutagens)?",
            "What do genes code for? If the genetic code changes, what happens to amino acids linked into a protein, and how can that alter the protein’s shape and its action in the cell? Explain the difference between a mutation that happens only in ordinary body tissue versus one that happens in germ-line cells—and which kind can offspring inherit?",
            "According to the natural-selection explanation in your notes, what must already exist inside a species before natural selection can act? Clarify whether the lesson says individual organisms evolve new traits within their lifetime or whether the population shifts across generations—and name the phrase the lesson uses when the environment favours certain traits.",
        ],
        "teacher_bullets": [
            "Adaptations intro: survival/reproduction passes traits; best-adapted survive; structural / behavioural / physiological definitions.",
            "Sorted example activity: classify listed examples against the three adaptation definitions.",
            "Sources of evolutionary change section: gradual change in traits over generations; variation from mutations; DNA-copy errors/radiation/mutagens cues.",
            "Mutation effects segment: genes → proteins → shape/action; somatic disappears with individual vs germ-line inherited.",
            "Natural selection recap: variation within species; population changes over generations; selective pressure from environment.",
        ],
        "diagram": (
            "Choose ONE sketch to draw and fully label from this list (diagrams only; labels on drawing):\n"
            "(i) One imaginary organism showing structural, behavioural, AND physiological adaptations (three labelled arrows or callouts).\n"
            "(ii) A simple labelled flow for natural selection—variation inside a species, environmental selective pressure or competition, surviving/reproducing individuals passing traits, and the population shifting over generations."
        ),
        "diagram_keys": [
            "(i) Structural = physical feature; behavioural = what it does to survive; physiological = internal/body-process adaptation tied to habitat.",
            "(ii) Flow should include variation requirement, selective pressure/environment, reproductive success transmitting traits, change at population scale.",
        ],
        "diagram_topic_alignment": "Adaptation categories OR natural-selection pathway (same-topic segments as homework questions).",
        "risk_notes": [
            "Opening/closing objectives pages repeat SWBAT wording—thin content.",
            "Image-heavy gag or cartoon frames near the adaptations unit may lack extractable text; accept student paraphrases.",
        ],
    },
    {
        "run_id": "bio20-mechanisms-lesson6-slides32-48-v2",
        "lesson": 6,
        "slide_span": "32-48",
        "topic": "Historical Development of Natural Selection Theory",
        "processed": "bio20-mechanisms-lesson6-slides32-48",
        "questions": [
            "Explain how theories about life’s history arise through scientific work—the lesson mentions observing nature, analysing data, and forming hypotheses—and name the overarching theory phrase the overview ends with. How did Plato and Aristotle view life forms on Earth (fixed and perfected versus changing)?",
            "What breakthroughs link Georges Cuvier to fossils in sedimentary strata, especially the pattern that fossils change as rock layers deepen? According to your notes, what kinds of catastrophes did his “revolutions” idea emphasize when species supposedly changed?",
            "How did Charles Lyell’s picture of geology differ from catastrophe-driven replacement—especially concerning how long geological processes take and whether slow, continual change applies to organisms as well?",
            "What misconception does your notepack emphasize about Lamarck’s “inheritance of acquired characteristics”—use the muscular-parent analogy to clarify what Lamarck wrongly claimed?",
            "Connect Thomas Malthus’s population pressure idea to Darwin and Wallace reasoning: why might competition for scarce resources favour some traits over others and change what proportion of a population expresses those traits over time? List the two major ideas the lesson attributes to Darwin’s book On the Origin of Species concerning ancestry and mechanism.",
        ],
        "teacher_bullets": [
            "Theory-development framing: observations/data/hypotheses evolve into evolution by natural selection; Greek fixed-perfection worldview.",
            "Cuvier: paleontology/stratigraphy storyline; fossils differ with depth; revolutions catastrophes replace species.",
            "Lyell contrasts with catastrophism; old Earth gradual processes extended metaphorically to life.",
            "Lamarck: acquired traits falsely believed inheritable.",
            "Malthus overcrowding offspring; Wallace/Darwin competition selects traits; proportional shift; Darwin book = descent modification + natural selection long timescale.",
        ],
        "diagram": (
            "Draw ONE labelled concept map tying together the lesson storyline: thinkers who questioned fixed species,"
            " Cuvier’s fossil-layer catastrophes, Lyell’s gradual deep-time geology,"
            " and Wallace/Darwin’s conclusion that environmental competition interacts with inherited variation via natural selection."
        ),
        "diagram_keys": [
            "Timeline/chronology arrows optional—must show causal links from fossil discoveries & Earth history thinkers to Wallace/Darwin synthesis.",
            "Key vocabulary checkpoints: paleontology, gradualism/stratigraphy contrasts, voyages observations, offspring competition mechanism.",
        ],
        "diagram_topic_alignment": "History-of-science storyline from Buffon-era challenges through Lyell/Darwin synthesis.",
        "risk_notes": [
            "Some biography slides are caption-light in extracted text—grade from class discussion if needed.",
            "Timeline graphics may differ between teacher deck vs notepack; accept equivalent sequencing.",
        ],
    },
    {
        "run_id": "bio20-mechanisms-lesson7-slides50-68-v4",
        "doc_version": "v4",
        "lesson": 7,
        "slide_span": "50-68",
        "topic": "Evidence for Evolution and Sources of Genetic Variability",
        "processed": "bio20-mechanisms-lesson7-slides50-68",
        "questions": [
            "After Darwin published On the Origin of Species, the lesson lists several scientific communities that kept gathering supporting evidence—which broad science fields contributed (beyond biologists)?",
            "Using the fossil-record arguments in your notes: how sedimentary rock hosts fossils; how layering hints at age order; why fossils nearer the surface resemble modern organisms more than fossils from deeper strata; how chronological layering supports ancestral vs younger species; how transitional fossils help close apparent gaps.",
            "Define biogeography and note why Wallace and Darwin relied on geography-based patterns. Explain how geographically close environments tend to harbour related organisms compared with distant places that merely look environmentally similar—and give the desert-cactus geography contrast plus one island-vs-mainland relatedness example cited in class.",
            "Define homologous structures versus analogous structures in the anatomical-evidence lesson. Explain how bird wings and bat wings illustrate analogy versus homology.",
            {
                "stem": (
                    "Your notepack’s evidence-for-evolution wrap-up ties together several modern arguments. "
                    "For each of the four areas below, write at least one complete sentence that (1) restates "
                    "the evidence in your own words and (2) explains how that evidence supports evolutionary "
                    "change or relatedness over time:"
                ),
                "subitems": [
                    "(a) embryology—many vertebrate embryos look very similar early in development;",
                    "(b) molecular biology—shared cell chemistry and enzymes, plus comparing DNA or genetic similarity across species;",
                    "(c) genetics since Darwin’s era—mutations as a source of inherited variation;",
                    "(d) sexual reproduction—offspring inherit two parental copies of each gene and can show new combinations of traits.",
                ],
            },
        ],
        "teacher_bullets": [
            "Multi-field evidence infographic: geology/paleontology/geography etc. supporting evolutionary theory.",
            "Fossils: sediment deposition; strata as age proxy; nearer-surface resemblance; ancestors deeper downward; transitional forms bridge gaps.",
            "Biogeography: distribution science; geographically close correlate relatedness; similar distant habitats can host unrelated fauna; Americas-only cactus vs Old-World desert plants; Canary lizards resembling West African relatives.",
            "Homology/common ancestry anatomy vs analogous same-function different origins wings example.",
            "Embryology/molecular/genetics/sexual-repro evidence cluster: each part needs restatement + explicit link to evolution (common ancestry, genetic change, or variability).",
        ],
        "diagram": (
            "Choose ONE diagram sketch to label fully:\n"
            "(A) Homologous forelimbs comparing at least three vertebrate organisms—same underlying bone blueprint, different lifestyles.\n"
            "(B) Side-by-side bird wing vs bat wing labelled to show analogous flight function but anatomically divergent evolutionary origins."
        ),
        "diagram_keys": [
            "(A) Label homologous structures + ancestral origin language.",
            "(B) Label analogous convergence + differing structural origins + shared flight ecological role.",
        ],
        "diagram_topic_alignment": "Comparative anatomy evidence unit (homology/analogy diagrams).",
        "risk_notes": [
            "Whale transitional sequence artwork may rely on captions not present in summaries—reward accurate paraphrases.",
            "Forelimb comparatives may compress labels in scanned notepacks—students may annotate generically.",
        ],
    },
    {
        "run_id": "bio20-mechanisms-lesson8-slides73-87-v2",
        "lesson": 8,
        "slide_span": "73-87",
        "topic": "Speciation, Isolation Patterns, Rates of Evolution",
        "processed": "bio20-mechanisms-lesson8-slides73-87",
        "questions": [
            "State the reproductive definition of a species emphasized in class (successful reproduction producing viable fertile offspring). Why might two populations look superficially similar yet count as distinct species?",
            "Describe geographic isolation versus behavioural isolation (include the mating-call illustration for behavioural isolation). Why does reproductive isolation keep species genetically separate?",
            "Explain transformation speciation versus divergence using the contrasts your notes emphasize—especially whether biological diversity rises or stays flat when transformation replaces older forms.",
            "How can long separation paired with mutations and divergent selection make populations incapable of mating even after a geographical barrier disappears? Then contrast allopatric speciation (physical separation examples like mountains or oceans) versus sympatric isolation living in overlapping ranges.",
            "Define adaptive radiation: diversification from one ancestor into specialised forms as descendant lineages exploit new habitats and face different selective pressures. Contrast gradualism with punctuated equilibrium—long stretches of comparative stability interrupted by comparatively rapid bursts of morphological change when lineages diverge or enter novel environments.",
        ],
        "teacher_bullets": [
            "Species concept: mating viability fertile offspring distinguishes lineages irrespective of morphology.",
            "Isolation mechanisms: geography vs mating behaviours/calls reinforcing separation.",
            "Transformation replaces ancestor (diversity nuanced) vs divergence adds species while parental lineage persists.",
            "Isolation duration accumulates divergence; reunion test of reproductive compatibility; geographical vs behavioural/chemical mating barriers enumerated in lesson.",
            "Adaptive radiation branching niches gradualism smoothing vs punctuation bursts stasis interplay.",
            "Graphic organizers/summary bullets near unit end recap common descent—all optional enrichment.",
        ],
        "diagram": (
            "Pick ONE labelled sketch:\n"
            "(1) Allopatric speciation storyline—ancestor split by geography into two accumulating differences.\n"
            "(2) Adaptive radiation branching from one ancestor outward with differing habitat/selective-pressure callouts.\n"
            "(3) Dual timelines contrasting gradual continual morphological drift versus punctuated long plateaus interrupted by bursts of rapid change."
        ),
        "diagram_keys": [
            "(1) Show barrier duration relevance + inability to reunite sexually after divergence.",
            "(2) Emphasise adaptive radiation diversification language + environmental drivers.",
            "(3) Contrast evolutionary tempo models gradualism versus punctuated equilibrium shapes.",
        ],
        "diagram_topic_alignment": "Speciation modes & evolutionary tempo summaries.",
        "risk_notes": [
            "Certain frames reference external enrichment videos—students should ignore hyperlink-only prompts unless taught.",
            "End-of-unit graphic organizer OCR may clash with notepack ordering—prioritize thematic accuracy over layout.",
        ],
    },
]


def _add_student_question_block(doc: Document, number: int, q: dict | str) -> None:
    """Emit student question text with native Word layout (no Markdown markup)."""

    if isinstance(q, str):
        doc.add_paragraph(f"{number}. {q}")
        doc.add_paragraph("")
        return

    stem = str(q.get("stem", "")).strip()
    doc.add_paragraph(f"{number}. {stem}")
    for line in q.get("subitems", []):
        p = doc.add_paragraph(line.strip())
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.22)
        p.paragraph_format.space_after = Pt(3)
    doc.add_paragraph("")


def set_body_font(doc: Document) -> None:
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)


def write_bundle(spec: dict) -> Path:
    ver = spec.get("doc_version", DOC_VERSION)
    outdir = OUTPUT_PARENT / spec["run_id"]
    outdir.mkdir(parents=True, exist_ok=True)
    stem = f"bio20-mechanisms-lesson{spec['lesson']}-{ver}"

    title = (
        f"Lesson {spec['lesson']} Homework — {spec['topic']} "
        f"(Mechanisms of Population Change)"
    )

    doc = Document()
    set_body_font(doc)
    p0 = doc.add_paragraph()
    r0 = p0.add_run(title)
    r0.bold = True
    r0.font.size = Pt(16)
    r0.font.name = "Arial"
    doc.add_paragraph(
        "Questions 1-5: answer in complete sentences. Question 6: diagrams and labels only on the diagrams page."
    )
    doc.add_paragraph(
        "Use your class notes and notepack; wording may follow the lesson storyline without matching deck order exactly."
    )
    doc.add_paragraph("")
    for i, q in enumerate(spec["questions"], 1):
        _add_student_question_block(doc, i, q)
    pb = doc.add_paragraph()
    pb.add_run().add_break(WD_BREAK.PAGE)
    h = doc.add_paragraph()
    hr = h.add_run("6. Diagram question")
    hr.bold = True
    doc.add_paragraph("")
    doc.add_paragraph(spec["diagram"])
    doc.save(outdir / f"{stem}-student.docx")

    key_doc = Document()
    set_body_font(key_doc)
    kt = key_doc.add_paragraph()
    kr = kt.add_run(f"{title} - Teacher Key ({ver})")
    kr.bold = True
    kr.font.size = Pt(16)
    key_doc.add_paragraph(
        f"Mapped from teacher deck slice slides {spec['slide_span']} (see processed JSON)."
    )
    for i, b in enumerate(spec["teacher_bullets"], 1):
        p = key_doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(b)
    p6 = key_doc.add_paragraph()
    p6.add_run("6. Diagram question - marking notes: ").bold = True
    for line in spec["diagram_keys"]:
        key_doc.add_paragraph(line)
    key_doc.add_paragraph("")
    key_doc.add_paragraph(
        f"Diagram alignment (internal): {spec['diagram_topic_alignment']} "
        f"within deck span {spec['slide_span']}."
    )

    risks = "; ".join(spec.get("risk_notes", []))
    if risks:
        key_doc.add_paragraph("")
        key_doc.add_paragraph(f"Risks / uncertainty: {risks}")

    key_doc.save(outdir / f"{stem}-teacher-key.docx")

    risk_notes = spec.get("risk_notes", [])
    risks_section = (
        "\n".join(f"- {n}" for n in risk_notes) if risk_notes else "- None noted"
    )
    audit = "\n".join(
        [
            f"# Lesson {spec['lesson']} homework audit ({ver})",
            "",
            "**Sourcedeck**: `2 - Mechanisms of Population Change (1).pptx`",
            f"**Teacher deck span (traceability only)**: {spec['slide_span']}",
            f"**Processed JSON**: `school-agents/data/processed/{spec['processed']}/slide-summaries.json`",
            f"**Run folder**: `{outdir.relative_to(ROOT)}`",
            "",
            "**Student-visible policy**: stems avoid slide/page numbers because notepacks and live decks diverge;",
            "questions anchor to topics and note language instead.",
            "",
            "## Diagram options",
            spec["diagram"],
            "",
            "### Topic alignment note",
            spec["diagram_topic_alignment"],
            "",
            "### Risks",
            risks_section,
        ]
    )
    (outdir / f"{stem}-audit.md").write_text(audit, encoding="utf-8")

    readme_lines = [
        f"# Lesson {spec['lesson']} Mechanisms homework ({ver})",
        "",
        f"- Student: `{stem}-student.docx` (no slide numbers in prompts)",
        f"- Teacher key: `{stem}-teacher-key.docx` (includes optional deck-span trace line)",
        f"- Audit / traceability: `{stem}-audit.md`",
        "",
        "Processed summaries for mapping: "
        f"`data/processed/{spec['processed']}/`.",
        "",
        "**v1 folders** retain earlier slide-numbered wording if needed for comparison.",
    ]
    (outdir / "README.md").write_text("\n".join(readme_lines) + "\n", encoding="utf-8")

    return outdir


def main() -> None:
    for spec in LESSONS:
        folder = write_bundle(spec)
        print(folder)


if __name__ == "__main__":
    main()
