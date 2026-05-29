"""
Build the Science 30 - Unit C (Electromagnetic Energy) review/test-prep DOCX.

Output:
- school-agents/outputs/homework/sci30-unitc-emr-review-v1/
    sci30-unitc-emr-review-student.docx  (combined student + GRAS(S) solutions)
    sci30-unitc-emr-review-summary.md    (template-aligned audit/summary)

Source authority:
- Alberta Science 30 data booklet pages 2-3 (gravitational/electric fields,
  astronomy data, electricity, waves).
- "Science 30 - Unit C - Electromagnetic Energy.pptx"
- "BONUS - Unit 3 Review Sheet.docx"

Question design:
- 10 equation types x 3 questions = 30 questions.
- Difficulty pattern per type: Straightforward, Straightforward, Challenging.
- Every question is a multi-step word problem.
- Space-themed contexts (Rosetta/Philae, Mars rover, ISS, Hubble, Cassini,
  lunar habitat, Voyager-style probes, etc.).
"""

from __future__ import annotations
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm

ROOT = Path(__file__).resolve().parent.parent
RUN_ID = "sci30-unitc-emr-review-v1"
OUT_DIR = ROOT / "outputs" / "homework" / RUN_ID
OUT_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT_DIR / "sci30-unitc-emr-review-student.docx"
MD_PATH = OUT_DIR / "sci30-unitc-emr-review-summary.md"


# --------------------------------------------------------------------------------------
# Question data model
# --------------------------------------------------------------------------------------


@dataclass
class Question:
    qid: str                     # e.g. "1.1"
    type_label: str              # e.g. "F_g = mg (Force due to gravity)"
    difficulty: str              # "Straightforward" or "Challenging"
    title: str                   # short scenario name
    prompt: str                  # full word-problem prompt (multi-step required)
    given: list[str]             # bulleted given list
    required: list[str]          # bulleted required list
    equation: list[str]          # bulleted equation list (from data booklet)
    solve: list[str]             # ordered solve steps (each step = a string line)
    statement: str               # final answer statement with units + sig figs
    final_answer: str            # short final answer summary
    annotated: list[str]         # annotated thinking bullets (what + why + common mistake)


# --------------------------------------------------------------------------------------
# Section banner: equation type
# --------------------------------------------------------------------------------------


@dataclass
class SectionType:
    code: str                    # e.g. "TYPE 1"
    name: str                    # e.g. "Force due to gravity"
    equations: list[str]         # equation strings as written on data sheet
    notes: str                   # short reminder about when to use
    questions: list[Question] = field(default_factory=list)


# --------------------------------------------------------------------------------------
# Build the question content
# --------------------------------------------------------------------------------------


def build_sections() -> list[SectionType]:
    sections: list[SectionType] = []

    # ===================================================================================
    # TYPE 1: F_g = mg
    # ===================================================================================
    s1 = SectionType(
        code="TYPE 1",
        name="Force due to gravity (F\u20D7\u2092 = mg\u20D7)",
        equations=["F\u20D7g = mg\u20D7"],
        notes=(
            "Use when you have a mass (kg) sitting in a gravitational field of known "
            "strength g (N/kg). Result is the weight (N) of the object at that location."
        ),
    )

    s1.questions.append(Question(
        qid="1.1",
        type_label="F_g = mg",
        difficulty="Straightforward",
        title="Philae lander on comet 67P",
        prompt=(
            "In November 2014, the Philae lander separated from the Rosetta orbiter and "
            "touched down on comet 67P/Churyumov\u2013Gerasimenko. The lander itself has a "
            "mass of 100. kg, and the science payload bolted to its frame adds another "
            "27 kg. At the touchdown site, the gravitational field strength of the comet "
            "is about 1.0 \u00D7 10\u207B\u00B3 N/kg. "
            "Calculate the total weight of the loaded lander on the comet\u2019s surface."
        ),
        given=[
            "m\u2081 (lander) = 100. kg",
            "m\u2082 (payload) = 27 kg",
            "g (at 67P touchdown site) = 1.0 \u00D7 10\u207B\u00B3 N/kg",
        ],
        required=["F\u20D7g (total weight on the comet) in N"],
        equation=["F\u20D7g = mg\u20D7"],
        solve=[
            "Step 1 \u2014 Find the total mass of the loaded lander:",
            "    m_total = m\u2081 + m\u2082 = 100. kg + 27 kg = 127 kg",
            "Step 2 \u2014 Apply F\u20D7g = mg\u20D7 with the comet\u2019s field strength:",
            "    F\u20D7g = (127 kg)(1.0 \u00D7 10\u207B\u00B3 N/kg)",
            "    F\u20D7g = 0.127 N",
            "Step 3 \u2014 Round to the smallest number of sig figs in the data (g has 2 s.f.):",
            "    F\u20D7g \u2248 0.13 N",
        ],
        statement="The loaded Philae lander weighs about 0.13 N on the surface of comet 67P.",
        final_answer="F\u20D7g \u2248 0.13 N (down, toward 67P\u2019s centre)",
        annotated=[
            "What: Add the two masses first because both are sitting in the same "
            "gravitational field, so weight depends on the total mass present.",
            "Why: F = mg is linear in m \u2014 doubling m doubles the weight. Forgetting the "
            "payload would underestimate the contact force on the comet.",
            "Common mistake: Plugging in g = 9.81 N/kg out of habit. On 67P the field is "
            "thousands of times weaker, which is why Philae bounced after touchdown.",
        ],
    ))

    s1.questions.append(Question(
        qid="1.2",
        type_label="F_g = mg",
        difficulty="Straightforward",
        title="Perseverance rover on Mars vs. on Earth",
        prompt=(
            "The Mars rover Perseverance has a mass of 1.025 \u00D7 10\u00B3 kg. On the Martian "
            "surface, the gravitational field strength is 3.71 N/kg. Determine "
            "(a) Perseverance\u2019s weight on Mars, "
            "(b) the rover\u2019s weight back on Earth\u2019s surface, and "
            "(c) the difference in weight between the two locations."
        ),
        given=[
            "m = 1.025 \u00D7 10\u00B3 kg",
            "g_Mars = 3.71 N/kg",
            "g_Earth = 9.81 N/kg (data booklet)",
        ],
        required=[
            "F\u20D7g on Mars",
            "F\u20D7g on Earth",
            "\u0394F\u20D7g (Earth weight \u2212 Mars weight)",
        ],
        equation=["F\u20D7g = mg\u20D7"],
        solve=[
            "(a) Weight on Mars:",
            "    F\u20D7g,Mars = (1.025 \u00D7 10\u00B3 kg)(3.71 N/kg)",
            "    F\u20D7g,Mars = 3.80 \u00D7 10\u00B3 N",
            "(b) Weight on Earth:",
            "    F\u20D7g,Earth = (1.025 \u00D7 10\u00B3 kg)(9.81 N/kg)",
            "    F\u20D7g,Earth = 1.01 \u00D7 10\u2074 N",
            "(c) Difference (how much \u201Clighter\u201D it feels on Mars):",
            "    \u0394F\u20D7g = 1.01 \u00D7 10\u2074 N \u2212 3.80 \u00D7 10\u00B3 N = 6.25 \u00D7 10\u00B3 N",
        ],
        statement=(
            "Perseverance weighs 3.80 \u00D7 10\u00B3 N on Mars and 1.01 \u00D7 10\u2074 N on Earth, "
            "a difference of 6.25 \u00D7 10\u00B3 N."
        ),
        final_answer=(
            "Mars: 3.80 \u00D7 10\u00B3 N; Earth: 1.01 \u00D7 10\u2074 N; \u0394 = 6.25 \u00D7 10\u00B3 N"
        ),
        annotated=[
            "What: Re-use F = mg with two different g values, then subtract.",
            "Why: Mass is a property of the object and does not change between planets. "
            "Only the field strength changes, so weight changes.",
            "Common mistake: Saying \u201Cthe rover loses mass on Mars.\u201D It does not \u2014 it "
            "loses weight. Always read what the question wants: mass (kg) or weight (N).",
        ],
    ))

    s1.questions.append(Question(
        qid="1.3",
        type_label="F_g = mg",
        difficulty="Challenging",
        title="EVA astronaut on Earth, Moon, and a low-g asteroid",
        prompt=(
            "An astronaut in a spacesuit and tool harness has a combined mass of 1.85 \u00D7 "
            "10\u00B2 kg. Calculate the astronaut\u2019s weight in three locations: "
            "(a) on Earth\u2019s surface, (b) on the Moon\u2019s surface where g = 1.62 N/kg, and "
            "(c) on the surface of asteroid Bennu where g = 7.8 \u00D7 10\u207B\u2075 N/kg. "
            "(d) Express the Moon weight as a percentage of the Earth weight, and "
            "comment briefly on why \u201Cwalking\u201D on Bennu is impossible."
        ),
        given=[
            "m = 1.85 \u00D7 10\u00B2 kg",
            "g_Earth = 9.81 N/kg",
            "g_Moon = 1.62 N/kg",
            "g_Bennu = 7.8 \u00D7 10\u207B\u2075 N/kg",
        ],
        required=[
            "F\u20D7g on Earth, Moon, Bennu",
            "Moon weight as a % of Earth weight",
        ],
        equation=["F\u20D7g = mg\u20D7"],
        solve=[
            "(a) F\u20D7g,Earth = (185 kg)(9.81 N/kg) = 1.81 \u00D7 10\u00B3 N",
            "(b) F\u20D7g,Moon  = (185 kg)(1.62 N/kg) = 3.00 \u00D7 10\u00B2 N",
            "(c) F\u20D7g,Bennu = (185 kg)(7.8 \u00D7 10\u207B\u2075 N/kg) = 1.4 \u00D7 10\u207B\u00B2 N",
            "(d) Percent of Earth weight on the Moon:",
            "    %  = (F\u20D7g,Moon / F\u20D7g,Earth) \u00D7 100",
            "    %  = (3.00 \u00D7 10\u00B2 N / 1.81 \u00D7 10\u00B3 N) \u00D7 100 \u2248 16.5 %",
            "    (This is also just g_Moon/g_Earth = 1.62/9.81 \u2248 0.165.)",
        ],
        statement=(
            "Earth weight \u2248 1.81 \u00D7 10\u00B3 N; Moon weight \u2248 3.00 \u00D7 10\u00B2 N (\u2248 16.5 % of "
            "Earth); Bennu weight \u2248 1.4 \u00D7 10\u207B\u00B2 N. On Bennu the weight is so small that "
            "any push (a footstep) easily exceeds it, so an astronaut would drift off the "
            "surface rather than walk."
        ),
        final_answer=(
            "Earth 1.81\u00D710\u00B3 N; Moon 3.00\u00D710\u00B2 N (16.5%); Bennu 1.4\u00D710\u207B\u00B2 N"
        ),
        annotated=[
            "What: Three direct F = mg calculations, then a ratio for the percentage.",
            "Why: Comparing weights on the same body is just comparing g values \u2014 the "
            "mass cancels in the ratio.",
            "Common mistake: Mixing up percentage of weight vs. percentage of mass. "
            "Always set up the ratio with units that match (N over N).",
        ],
    ))
    sections.append(s1)

    # ===================================================================================
    # TYPE 2: g = Gm/r^2
    # ===================================================================================
    s2 = SectionType(
        code="TYPE 2",
        name="Gravitational field strength at a distance (g = Gm/r\u00B2)",
        equations=["g = Gm / r\u00B2"],
        notes=(
            "Use to find the field strength produced by a single source mass m at a "
            "distance r from its centre. Always use centre-to-centre distance, not "
            "altitude-above-the-surface."
        ),
    )

    s2.questions.append(Question(
        qid="2.1",
        type_label="g = Gm/r\u00B2",
        difficulty="Straightforward",
        title="Field strength at the ISS orbit",
        prompt=(
            "The International Space Station orbits about 4.20 \u00D7 10\u2075 m above Earth\u2019s "
            "surface. Treating Earth as a point mass at its centre, calculate the "
            "gravitational field strength on the ISS at orbital altitude, and compare it "
            "to the value at Earth\u2019s surface from the data booklet."
        ),
        given=[
            "altitude h = 4.20 \u00D7 10\u2075 m above the surface",
            "M_Earth = 5.98 \u00D7 10\u00B2\u2074 kg",
            "R_Earth = 6.37 \u00D7 10\u2076 m",
            "G = 6.67 \u00D7 10\u207B\u00B9\u00B9 N\u00B7m\u00B2/kg\u00B2",
        ],
        required=["g at the ISS orbit"],
        equation=["g = Gm / r\u00B2"],
        solve=[
            "Step 1 \u2014 Convert altitude to centre-to-centre distance:",
            "    r = R_Earth + h = 6.37 \u00D7 10\u2076 m + 4.20 \u00D7 10\u2075 m = 6.79 \u00D7 10\u2076 m",
            "Step 2 \u2014 Substitute into g = GM/r\u00B2:",
            "    g = (6.67 \u00D7 10\u207B\u00B9\u00B9)(5.98 \u00D7 10\u00B2\u2074) / (6.79 \u00D7 10\u2076)\u00B2",
            "    g = (3.989 \u00D7 10\u00B9\u2074) / (4.610 \u00D7 10\u00B9\u00B3)",
            "    g = 8.65 N/kg",
            "Step 3 \u2014 Compare to surface (9.81 N/kg): the ISS is still in a fairly strong "
            "field. Astronauts feel \u2018weightless\u2019 because they are in continuous free fall, "
            "not because gravity is gone.",
        ],
        statement=(
            "The gravitational field strength at the ISS orbit is g \u2248 8.65 N/kg, only "
            "about 12 % weaker than at Earth\u2019s surface."
        ),
        final_answer="g \u2248 8.65 N/kg",
        annotated=[
            "What: Add Earth\u2019s radius to altitude before squaring \u2014 the equation needs "
            "centre-to-centre distance.",
            "Why: Field lines emanate from the centre of mass; the surface is just one "
            "specific r-value from that centre.",
            "Common mistake: Using only the altitude (4.20 \u00D7 10\u2075 m) for r. That would "
            "predict a field hundreds of times too strong.",
        ],
    ))

    s2.questions.append(Question(
        qid="2.2",
        type_label="g = Gm/r\u00B2",
        difficulty="Straightforward",
        title="Sun\u2019s gravitational field at Earth and at half-distance",
        prompt=(
            "Earth orbits the Sun at an average distance of 1 AU. Imagine a future "
            "spacecraft using a solar sail to fly inward to a parking orbit at 0.500 AU. "
            "(a) Determine the Sun\u2019s gravitational field strength at 1.00 AU. "
            "(b) Determine the field strength at 0.500 AU. "
            "(c) By what factor is the field at 0.500 AU stronger than the field at 1.00 AU, "
            "and does this match the inverse-square prediction?"
        ),
        given=[
            "M_Sun = 1.99 \u00D7 10\u00B3\u2070 kg",
            "1 AU = 1.50 \u00D7 10\u00B9\u00B9 m",
            "G = 6.67 \u00D7 10\u207B\u00B9\u00B9 N\u00B7m\u00B2/kg\u00B2",
        ],
        required=["g at 1.00 AU", "g at 0.500 AU", "ratio g(0.5 AU)/g(1 AU)"],
        equation=["g = Gm / r\u00B2"],
        solve=[
            "(a) At r\u2081 = 1.50 \u00D7 10\u00B9\u00B9 m:",
            "    g\u2081 = (6.67 \u00D7 10\u207B\u00B9\u00B9)(1.99 \u00D7 10\u00B3\u2070) / (1.50 \u00D7 10\u00B9\u00B9)\u00B2",
            "    g\u2081 = (1.327 \u00D7 10\u00B2\u2070) / (2.25 \u00D7 10\u00B2\u00B2)",
            "    g\u2081 = 5.90 \u00D7 10\u207B\u00B3 N/kg",
            "(b) At r\u2082 = 0.500 AU = 7.50 \u00D7 10\u00B9\u2070 m:",
            "    g\u2082 = (1.327 \u00D7 10\u00B2\u2070) / (7.50 \u00D7 10\u00B9\u2070)\u00B2",
            "    g\u2082 = (1.327 \u00D7 10\u00B2\u2070) / (5.625 \u00D7 10\u00B2\u00B9)",
            "    g\u2082 = 2.36 \u00D7 10\u207B\u00B2 N/kg",
            "(c) Ratio:",
            "    g\u2082 / g\u2081 = (2.36 \u00D7 10\u207B\u00B2) / (5.90 \u00D7 10\u207B\u00B3) = 4.00",
            "    Inverse-square prediction: halving r should multiply g by (1/0.5)\u00B2 = 4. \u2713",
        ],
        statement=(
            "g(1 AU) \u2248 5.90 \u00D7 10\u207B\u00B3 N/kg and g(0.5 AU) \u2248 2.36 \u00D7 10\u207B\u00B2 N/kg. "
            "Halving the distance produces a 4-fold increase in field strength, "
            "matching the inverse-square law exactly."
        ),
        final_answer=(
            "g(1 AU) = 5.90\u00D710\u207B\u00B3 N/kg; g(0.5 AU) = 2.36\u00D710\u207B\u00B2 N/kg; ratio = 4.00"
        ),
        annotated=[
            "What: Same equation, two distances, then a ratio.",
            "Why: The factor-of-4 result is the signature of a 1/r\u00B2 field. Recognising "
            "this is faster than computing in some MC questions.",
            "Common mistake: Using AU directly in the equation. Convert to metres first, "
            "or your units will not work out to N/kg.",
        ],
    ))

    s2.questions.append(Question(
        qid="2.3",
        type_label="g = Gm/r\u00B2",
        difficulty="Challenging",
        title="Rosetta orbiting comet 67P (combines g = Gm/r\u00B2 with F = mg)",
        prompt=(
            "The Rosetta spacecraft (mass 1.30 \u00D7 10\u00B3 kg, including remaining fuel) "
            "orbited comet 67P at a distance of 30.0 km from the comet\u2019s centre. The "
            "comet\u2019s mass is 1.0 \u00D7 10\u00B9\u00B3 kg. "
            "(a) Calculate the gravitational field strength produced by 67P at Rosetta\u2019s "
            "orbital location. "
            "(b) Calculate the gravitational force the comet exerts on Rosetta at that "
            "distance. "
            "(c) Briefly explain why this orbit needed careful navigation, given the size "
            "of the force you calculated."
        ),
        given=[
            "m_67P = 1.0 \u00D7 10\u00B9\u00B3 kg",
            "r = 30.0 km = 3.00 \u00D7 10\u2074 m (centre-to-centre)",
            "m_Rosetta = 1.30 \u00D7 10\u00B3 kg",
            "G = 6.67 \u00D7 10\u207B\u00B9\u00B9 N\u00B7m\u00B2/kg\u00B2",
        ],
        required=["g at Rosetta\u2019s orbit", "F\u20D7g on Rosetta"],
        equation=["g = Gm / r\u00B2", "F\u20D7g = mg\u20D7"],
        solve=[
            "(a) Field strength at the orbit:",
            "    g = (6.67 \u00D7 10\u207B\u00B9\u00B9)(1.0 \u00D7 10\u00B9\u00B3) / (3.00 \u00D7 10\u2074)\u00B2",
            "    g = (6.67 \u00D7 10\u00B2) / (9.00 \u00D7 10\u2078)",
            "    g = 7.4 \u00D7 10\u207B\u2077 N/kg",
            "(b) Force on Rosetta:",
            "    F\u20D7g = mg\u20D7 = (1.30 \u00D7 10\u00B3 kg)(7.4 \u00D7 10\u207B\u2077 N/kg)",
            "    F\u20D7g = 9.6 \u00D7 10\u207B\u2074 N",
            "(c) The pull is less than a milligram of weight on Earth. Tiny "
            "manoeuvring thrusts, the pressure from solar wind, and even outgassing from "
            "the comet are comparable, so the orbit is fragile and must be re-tuned often.",
        ],
        statement=(
            "Rosetta sits in a field of about 7.4 \u00D7 10\u207B\u2077 N/kg around 67P, and feels a "
            "gravitational pull of only 9.6 \u00D7 10\u207B\u2074 N. The orbit is therefore extremely "
            "delicate \u2014 small disturbances easily perturb the trajectory."
        ),
        final_answer="g \u2248 7.4\u00D710\u207B\u2077 N/kg; F\u20D7g \u2248 9.6\u00D710\u207B\u2074 N",
        annotated=[
            "What: Compute g at the orbit first, then drop g into F = mg with the "
            "spacecraft\u2019s mass.",
            "Why: g = Gm/r\u00B2 only depends on the source (comet), not the test body "
            "(Rosetta). Once g is found at that location, any mass placed there feels "
            "F = mg.",
            "Common mistake: Confusing the two masses. The m in g = Gm/r\u00B2 is the source "
            "(67P); the m in F = mg is the test body (Rosetta).",
        ],
    ))
    sections.append(s2)

    # ===================================================================================
    # TYPE 3: |E| = kq/r^2
    # ===================================================================================
    s3 = SectionType(
        code="TYPE 3",
        name="Electric field strength (|E\u20D7| = kq/r\u00B2)",
        equations=["|E\u20D7| = kq / r\u00B2"],
        notes=(
            "Use for the field strength (in N/C) around a single point or sphere of "
            "charge q at a distance r. Direction: away from + charges, toward \u2212 charges."
        ),
    )

    s3.questions.append(Question(
        qid="3.1",
        type_label="|E\u20D7| = kq/r\u00B2",
        difficulty="Straightforward",
        title="Charged dust grain on the Moon",
        prompt=(
            "Lunar dust becomes electrically charged from solar UV. A single grain near "
            "an Apollo-style base picks up a charge of +5.00 nC. Find the electric field "
            "strength at a point 0.500 m away from the grain. Then state the direction "
            "of the field at that point."
        ),
        given=[
            "q = +5.00 nC = +5.00 \u00D7 10\u207B\u2079 C",
            "r = 0.500 m",
            "k = 8.99 \u00D7 10\u2079 N\u00B7m\u00B2/C\u00B2",
        ],
        required=["|E\u20D7| at the point", "direction of E\u20D7"],
        equation=["|E\u20D7| = kq / r\u00B2"],
        solve=[
            "Step 1 \u2014 Convert nC to C:",
            "    q = 5.00 nC \u00D7 (1 \u00D7 10\u207B\u2079 C / nC) = 5.00 \u00D7 10\u207B\u2079 C",
            "Step 2 \u2014 Substitute into |E| = kq/r\u00B2:",
            "    |E| = (8.99 \u00D7 10\u2079)(5.00 \u00D7 10\u207B\u2079) / (0.500)\u00B2",
            "    |E| = (44.95) / (0.250)",
            "    |E| = 1.80 \u00D7 10\u00B2 N/C",
            "Step 3 \u2014 Direction: q is positive, so E\u20D7 points away from the grain at "
            "every point in space.",
        ],
        statement=(
            "The electric field 0.500 m from the dust grain has magnitude "
            "|E\u20D7| \u2248 1.80 \u00D7 10\u00B2 N/C, pointing radially away from the grain."
        ),
        final_answer="|E\u20D7| \u2248 1.80 \u00D7 10\u00B2 N/C, directed away from the +q",
        annotated=[
            "What: Two steps \u2014 unit conversion, then plug-and-chug.",
            "Why: SI unit consistency matters. Forgetting the nC \u2192 C conversion gives an "
            "answer that is 10\u2079 times too large.",
            "Common mistake: Dropping the sign of q. The magnitude formula uses |q|, but "
            "the sign tells you the direction of the field.",
        ],
    ))

    s3.questions.append(Question(
        qid="3.2",
        type_label="|E\u20D7| = kq/r\u00B2",
        difficulty="Straightforward",
        title="Static-charged astronaut tool",
        prompt=(
            "While testing a tool inside a vacuum chamber on the ISS, a wrench picks up a "
            "static charge of \u22128.0 \u00D7 10\u207B\u2075 C. Determine the magnitude of the "
            "electric field this tool produces at (a) 0.10 m and (b) 0.30 m from its "
            "centre. (c) Comment on what doubling the distance would do to the field."
        ),
        given=[
            "q = \u22128.0 \u00D7 10\u207B\u2075 C  (sign indicates excess electrons)",
            "r\u2081 = 0.10 m, r\u2082 = 0.30 m",
            "k = 8.99 \u00D7 10\u2079 N\u00B7m\u00B2/C\u00B2",
        ],
        required=["|E\u20D7| at 0.10 m and 0.30 m"],
        equation=["|E\u20D7| = kq / r\u00B2"],
        solve=[
            "(a) At r\u2081 = 0.10 m:",
            "    |E\u2081| = (8.99 \u00D7 10\u2079)(8.0 \u00D7 10\u207B\u2075) / (0.10)\u00B2",
            "    |E\u2081| = (7.19 \u00D7 10\u2075) / (1.0 \u00D7 10\u207B\u00B2)",
            "    |E\u2081| = 7.2 \u00D7 10\u2077 N/C  (toward the wrench, since q is negative)",
            "(b) At r\u2082 = 0.30 m:",
            "    |E\u2082| = (8.99 \u00D7 10\u2079)(8.0 \u00D7 10\u207B\u2075) / (0.30)\u00B2",
            "    |E\u2082| = (7.19 \u00D7 10\u2075) / (9.0 \u00D7 10\u207B\u00B2)",
            "    |E\u2082| = 8.0 \u00D7 10\u2076 N/C",
            "(c) Tripling r divides the field by 3\u00B2 = 9 (\u22487.2\u00D710\u2077 / 9 \u2248 8\u00D710\u2076 \u2713). "
            "Doubling r would divide the field by 4.",
        ],
        statement=(
            "|E\u20D7| \u2248 7.2 \u00D7 10\u2077 N/C at 0.10 m and \u2248 8.0 \u00D7 10\u2076 N/C at 0.30 m. The ratio "
            "(0.30/0.10)\u00B2 = 9 matches the inverse-square drop in field strength."
        ),
        final_answer="|E\u20D7| \u2248 7.2\u00D710\u2077 N/C at 0.10 m; \u2248 8.0\u00D710\u2076 N/C at 0.30 m",
        annotated=[
            "What: Use the magnitude of q (drop the sign for size of E), then check using "
            "the inverse-square shortcut.",
            "Why: The sign of q only sets direction. For magnitude, |q| is enough.",
            "Common mistake: Squaring r once (writing /r) instead of /r\u00B2. Always double-"
            "check the exponent before pressing equals.",
        ],
    ))

    s3.questions.append(Question(
        qid="3.3",
        type_label="|E\u20D7| = kq/r\u00B2",
        difficulty="Challenging",
        title="Charged communications satellite",
        prompt=(
            "A geostationary communications satellite slowly accumulates a charge of "
            "+2.50 \u00D7 10\u207B\u2074 C from solar wind ions over many months. "
            "(a) Calculate the electric field strength produced by the satellite at 100. m "
            "from its centre. "
            "(b) Calculate the field strength at 300. m. "
            "(c) Determine the distance from the centre at which |E\u20D7| has dropped to "
            "1.00 \u00D7 10\u00B2 N/C. (d) Confirm the inverse-square relationship between your "
            "answers in (a) and (b)."
        ),
        given=[
            "q = +2.50 \u00D7 10\u207B\u2074 C",
            "k = 8.99 \u00D7 10\u2079 N\u00B7m\u00B2/C\u00B2",
            "r\u2081 = 100. m, r\u2082 = 300. m",
            "Target |E\u20D7| = 1.00 \u00D7 10\u00B2 N/C",
        ],
        required=[
            "|E\u20D7| at 100. m",
            "|E\u20D7| at 300. m",
            "r at which |E\u20D7| = 1.00 \u00D7 10\u00B2 N/C",
        ],
        equation=["|E\u20D7| = kq / r\u00B2", "rearranged: r = \u221A(kq / |E\u20D7|)"],
        solve=[
            "(a) At r\u2081 = 100. m:",
            "    |E\u2081| = (8.99 \u00D7 10\u2079)(2.50 \u00D7 10\u207B\u2074) / (100.)\u00B2",
            "    |E\u2081| = (2.248 \u00D7 10\u2076) / (1.00 \u00D7 10\u2074)",
            "    |E\u2081| = 2.25 \u00D7 10\u00B2 N/C",
            "(b) At r\u2082 = 300. m:",
            "    |E\u2082| = (2.248 \u00D7 10\u2076) / (9.00 \u00D7 10\u2074)",
            "    |E\u2082| = 25.0 N/C",
            "(c) Solve r from |E| = kq/r\u00B2:",
            "    r\u00B2 = kq / |E| = (8.99 \u00D7 10\u2079)(2.50 \u00D7 10\u207B\u2074) / (1.00 \u00D7 10\u00B2)",
            "    r\u00B2 = 2.248 \u00D7 10\u2074 m\u00B2",
            "    r = 1.50 \u00D7 10\u00B2 m  (i.e. 150 m)",
            "(d) Inverse-square check: r tripled (100 \u2192 300), so |E| should fall by 9. "
            "    225 / 25 = 9 \u2713",
        ],
        statement=(
            "|E\u20D7|(100 m) \u2248 2.25 \u00D7 10\u00B2 N/C; |E\u20D7|(300 m) \u2248 25.0 N/C; |E\u20D7| reaches "
            "1.00 \u00D7 10\u00B2 N/C at r \u2248 1.50 \u00D7 10\u00B2 m. The 1/r\u00B2 relationship is verified."
        ),
        final_answer="225 N/C; 25.0 N/C; 1.50\u00D710\u00B2 m; ratio 9 \u2713",
        annotated=[
            "What: Two direct E calculations, one solve-for-r rearrangement, and a check.",
            "Why: Practising rearranging r = \u221A(kq/|E|) is essential for diploma-style "
            "questions that hide which variable is the unknown.",
            "Common mistake: Forgetting the square root when solving for r. Always finish "
            "with units: r should come out in metres.",
        ],
    ))
    sections.append(s3)

    # ===================================================================================
    # TYPE 4: V = IR
    # ===================================================================================
    s4 = SectionType(
        code="TYPE 4",
        name="Ohm\u2019s Law (V = IR)",
        equations=["V = IR"],
        notes=(
            "Use whenever you have any two of voltage (V), current (A), and resistance "
            "(\u03A9) and need the third. Always use SI units."
        ),
    )

    s4.questions.append(Question(
        qid="4.1",
        type_label="V = IR",
        difficulty="Straightforward",
        title="Indicator lamp on a satellite bus",
        prompt=(
            "An indicator lamp on a small satellite operates from the 12.0-V power bus and "
            "draws 0.50 A while lit. "
            "(a) Determine the resistance of the lamp while it operates. "
            "(b) During a fault test, the bus voltage is doubled to 24.0 V. Assuming the "
            "lamp\u2019s resistance does not change, determine the new current."
        ),
        given=[
            "V\u2081 = 12.0 V, I\u2081 = 0.50 A",
            "V\u2082 = 24.0 V (after fault test)",
        ],
        required=["R", "I\u2082 at 24.0 V"],
        equation=["V = IR  \u2192  R = V/I,  I = V/R"],
        solve=[
            "(a) R = V\u2081 / I\u2081 = 12.0 V / 0.50 A = 24 \u03A9",
            "(b) I\u2082 = V\u2082 / R = 24.0 V / 24 \u03A9 = 1.0 A",
        ],
        statement=(
            "The lamp has a resistance of 24 \u03A9. When the bus voltage doubles to 24.0 V "
            "(at constant resistance), the current also doubles to 1.0 A."
        ),
        final_answer="R = 24 \u03A9; I\u2082 = 1.0 A",
        annotated=[
            "What: Apply Ohm\u2019s Law twice \u2014 once to find R, once to find new I.",
            "Why: At constant R, V and I are directly proportional. Doubling V doubles I.",
            "Common mistake: Using the original current 0.50 A in part (b). The current "
            "is what changes; only R is held fixed.",
        ],
    ))

    s4.questions.append(Question(
        qid="4.2",
        type_label="V = IR",
        difficulty="Straightforward",
        title="Mars rover heater",
        prompt=(
            "A small heater in the Mars rover\u2019s science instrument bay is rated at 25.0 "
            "\u03A9 of resistance and is powered by the rover\u2019s 12.0-V battery. "
            "(a) Determine the current through the heater. "
            "(b) After dust accumulates on the heating element, the effective resistance "
            "rises to 40.0 \u03A9. Determine the new current through the heater."
        ),
        given=[
            "R\u2081 = 25.0 \u03A9, V = 12.0 V",
            "R\u2082 = 40.0 \u03A9 (after dust)",
        ],
        required=["I\u2081 (clean)", "I\u2082 (dusty)"],
        equation=["V = IR  \u2192  I = V/R"],
        solve=[
            "(a) I\u2081 = V / R\u2081 = 12.0 V / 25.0 \u03A9 = 0.480 A",
            "(b) I\u2082 = V / R\u2082 = 12.0 V / 40.0 \u03A9 = 0.300 A",
        ],
        statement=(
            "The clean heater draws 0.480 A. With dust raising the resistance to 40.0 \u03A9, "
            "the current drops to 0.300 A \u2014 reducing the heating power."
        ),
        final_answer="I\u2081 = 0.480 A; I\u2082 = 0.300 A",
        annotated=[
            "What: Same equation, two scenarios, two values of R.",
            "Why: At fixed V, current and resistance are inversely related: more "
            "resistance means less current.",
            "Common mistake: Mixing up which variable is constant. Read the question to "
            "see whether V, I, or R is the one that changes.",
        ],
    ))

    s4.questions.append(Question(
        qid="4.3",
        type_label="V = IR",
        difficulty="Challenging",
        title="Damaged heater on the ISS",
        prompt=(
            "A heating element on the ISS originally draws 5.0 A from the 28-V utility "
            "bus. After insulation around the wire degrades, the element\u2019s resistance "
            "drops to half of its original value while the bus voltage stays at 28 V. "
            "(a) Calculate the original resistance of the heater. "
            "(b) Calculate the new resistance and the new current. "
            "(c) Briefly explain why this is a fire-hazard concern, even though the heater "
            "is still working."
        ),
        given=[
            "V = 28 V (constant)",
            "I_original = 5.0 A",
            "R_new = (1/2) R_original",
        ],
        required=["R_original", "R_new", "I_new"],
        equation=["V = IR"],
        solve=[
            "(a) Original resistance:",
            "    R_original = V / I_original = 28 V / 5.0 A = 5.6 \u03A9",
            "(b) Damaged resistance and current:",
            "    R_new = 0.50 \u00D7 5.6 \u03A9 = 2.8 \u03A9",
            "    I_new = V / R_new = 28 V / 2.8 \u03A9 = 1.0 \u00D7 10\u00B9 A   (= 10. A)",
            "(c) Halving R doubles the current, and (looking ahead at P = I\u00B2R) the heat "
            "produced rises sharply. Wires not rated for 10 A will overheat \u2014 a hazard.",
        ],
        statement=(
            "Original resistance was 5.6 \u03A9 with 5.0 A flowing. After the fault, "
            "R_new \u2248 2.8 \u03A9 and the current rises to 1.0 \u00D7 10\u00B9 A \u2014 enough to overheat "
            "wiring not rated for that load."
        ),
        final_answer="R_original = 5.6 \u03A9; R_new = 2.8 \u03A9; I_new = 1.0\u00D710\u00B9 A",
        annotated=[
            "What: Reverse Ohm\u2019s Law to get original R, halve R, then forward-solve for "
            "the new current.",
            "Why: Faults in real circuits often look like sudden current jumps \u2014 V is "
            "fixed by the supply, R drops, so I climbs.",
            "Common mistake: Halving the current instead of the resistance. Re-read the "
            "question and identify which quantity is changing.",
        ],
    ))
    sections.append(s4)

    # ===================================================================================
    # TYPE 5: Series resistance
    # ===================================================================================
    s5 = SectionType(
        code="TYPE 5",
        name="Series resistance (R_T = R\u2081 + R\u2082 + ... + R_n)",
        equations=["R_T = R\u2081 + R\u2082 + ... + R_n"],
        notes=(
            "Use when components share the same single current path. Add resistances; "
            "current is the same through every component; voltages add to the source."
        ),
    )

    s5.questions.append(Question(
        qid="5.1",
        type_label="R_T (series)",
        difficulty="Straightforward",
        title="String of LEDs in a CubeSat",
        prompt=(
            "A small CubeSat science indicator panel has four identical LED sub-modules "
            "wired in series, each with a resistance of 30. \u03A9. The panel is supplied by "
            "a 12.0-V regulator. Determine the total resistance of the string and the "
            "current that flows through each module."
        ),
        given=[
            "n = 4 modules in series",
            "R_each = 30. \u03A9",
            "V = 12.0 V",
        ],
        required=["R_T", "I (same through all modules)"],
        equation=["R_T = R\u2081 + R\u2082 + R\u2083 + R\u2084", "V = IR"],
        solve=[
            "Step 1 \u2014 Add the resistances:",
            "    R_T = 4 \u00D7 30. \u03A9 = 1.2 \u00D7 10\u00B2 \u03A9",
            "Step 2 \u2014 Apply V = IR for the whole loop:",
            "    I = V / R_T = 12.0 V / 1.2 \u00D7 10\u00B2 \u03A9 = 0.10 A",
        ],
        statement=(
            "The total resistance is R_T = 1.2 \u00D7 10\u00B2 \u03A9 and the current through every "
            "LED in the string is 0.10 A."
        ),
        final_answer="R_T = 1.2\u00D710\u00B2 \u03A9; I = 0.10 A",
        annotated=[
            "What: Add up identical resistances, then use Ohm\u2019s Law on the whole loop.",
            "Why: In series, the current is the same everywhere \u2014 you only need one I.",
            "Common mistake: Dividing the supply voltage by one resistor\u2019s value instead "
            "of by R_T. That gives the current you would have if only one LED were "
            "wired up, not four.",
        ],
    ))

    s5.questions.append(Question(
        qid="5.2",
        type_label="R_T (series)",
        difficulty="Straightforward",
        title="Series \u201Cstring\u201D lights inside a Mars habitat",
        prompt=(
            "Inside a simulated Mars habitat module, eight decorative bulbs are wired in "
            "series along one strand. Each bulb has a resistance of 64.0 \u03A9. The strand is "
            "plugged into a 24.0-V outlet. "
            "(a) Determine the total resistance of the strand. "
            "(b) Determine the current through any one bulb."
        ),
        given=[
            "n = 8 bulbs in series",
            "R_each = 64.0 \u03A9",
            "V = 24.0 V",
        ],
        required=["R_T", "I"],
        equation=["R_T = R\u2081 + R\u2082 + ... + R\u2088", "V = IR"],
        solve=[
            "(a) R_T = 8 \u00D7 64.0 \u03A9 = 512 \u03A9",
            "(b) I = V / R_T = 24.0 V / 512 \u03A9 = 4.69 \u00D7 10\u207B\u00B2 A   (\u2248 0.0469 A)",
        ],
        statement=(
            "The strand has a total resistance of 512 \u03A9. The current through any one "
            "bulb in the series strand is 4.69 \u00D7 10\u207B\u00B2 A."
        ),
        final_answer="R_T = 512 \u03A9; I = 4.69\u00D710\u207B\u00B2 A",
        annotated=[
            "What: Same approach as 5.1 with eight resistors instead of four.",
            "Why: Sum behaviour scales linearly. The bigger the string, the more "
            "resistance, the smaller the shared current.",
            "Common mistake: Saying \u201Cthe current splits among the bulbs.\u201D Current does "
            "not split in series \u2014 only voltages add up across components.",
        ],
    ))

    s5.questions.append(Question(
        qid="5.3",
        type_label="R_T (series)",
        difficulty="Challenging",
        title="Sensor train on a deep-space probe",
        prompt=(
            "A deep-space probe carries a sensor train wired in series along a single "
            "data line: a thermistor (20.0 \u03A9), a Hall-effect sensor (50.0 \u03A9), and a "
            "radiation detector (80.0 \u03A9). The probe\u2019s on-board regulator supplies 30.0 V "
            "to this train. "
            "(a) Determine the total resistance. "
            "(b) Determine the current flowing in the loop. "
            "(c) Determine the voltage drop across each individual sensor. "
            "(d) Verify that the voltage drops add up to the regulator voltage."
        ),
        given=[
            "R\u2081 = 20.0 \u03A9 (thermistor)",
            "R\u2082 = 50.0 \u03A9 (Hall-effect)",
            "R\u2083 = 80.0 \u03A9 (radiation detector)",
            "V = 30.0 V",
        ],
        required=["R_T", "I", "V across each sensor", "verify \u03A3V = source"],
        equation=["R_T = R\u2081 + R\u2082 + R\u2083", "V = IR (across each component)"],
        solve=[
            "(a) R_T = 20.0 + 50.0 + 80.0 = 150.0 \u03A9",
            "(b) I = V / R_T = 30.0 V / 150.0 \u03A9 = 0.200 A",
            "(c) Voltage drops across each sensor (current is the same):",
            "    V\u2081 = IR\u2081 = (0.200)(20.0) = 4.00 V",
            "    V\u2082 = IR\u2082 = (0.200)(50.0) = 10.0 V",
            "    V\u2083 = IR\u2083 = (0.200)(80.0) = 16.0 V",
            "(d) Sum: V\u2081 + V\u2082 + V\u2083 = 4.00 + 10.0 + 16.0 = 30.0 V \u2713",
        ],
        statement=(
            "R_T = 150.0 \u03A9; I = 0.200 A; voltage drops are 4.00 V, 10.0 V, and 16.0 V. "
            "These sum to 30.0 V, matching the regulator voltage \u2014 a direct check on the "
            "series result."
        ),
        final_answer="R_T = 150.0 \u03A9; I = 0.200 A; V\u2081/V\u2082/V\u2083 = 4.00/10.0/16.0 V (sum 30.0 V \u2713)",
        annotated=[
            "What: Sum resistances, find the shared current, then use V = IR component "
            "by component.",
            "Why: In a series circuit, the resistor with the largest R takes the largest "
            "share of the supply voltage \u2014 a useful diploma-exam shortcut.",
            "Common mistake: Computing each component\u2019s V using V = source/n. That only "
            "works if all resistors are equal. With unequal R, you must use V = IR.",
        ],
    ))
    sections.append(s5)

    # ===================================================================================
    # TYPE 6: Parallel resistance
    # ===================================================================================
    s6 = SectionType(
        code="TYPE 6",
        name="Parallel resistance (1/R_T = 1/R\u2081 + 1/R\u2082 + ... + 1/R_n)",
        equations=["1/R_T = 1/R\u2081 + 1/R\u2082 + ... + 1/R_n"],
        notes=(
            "Use when components share the same two nodes (each has the full source "
            "voltage across it). Total current splits between branches; total resistance "
            "is always smaller than the smallest branch."
        ),
    )

    s6.questions.append(Question(
        qid="6.1",
        type_label="R_T (parallel)",
        difficulty="Straightforward",
        title="Two parallel cabin lights on a Mars rover",
        prompt=(
            "Two identical 100.-\u03A9 cabin lights on a Mars rover are wired in parallel "
            "across the 24.0-V battery bus. "
            "(a) Determine the total resistance of the pair. "
            "(b) Determine the total current drawn from the battery."
        ),
        given=[
            "R\u2081 = R\u2082 = 100. \u03A9 (parallel)",
            "V = 24.0 V",
        ],
        required=["R_T", "I_total"],
        equation=[
            "1/R_T = 1/R\u2081 + 1/R\u2082",
            "V = IR  \u2192  I_total = V / R_T",
        ],
        solve=[
            "(a) 1/R_T = 1/100. + 1/100. = 2/100. = 1/50.",
            "    R_T = 50. \u03A9",
            "(b) I_total = V / R_T = 24.0 V / 50. \u03A9 = 0.480 A",
        ],
        statement=(
            "The pair behaves like a single 50.-\u03A9 resistor and pulls 0.480 A from the "
            "battery."
        ),
        final_answer="R_T = 50. \u03A9; I_total = 0.480 A",
        annotated=[
            "What: Add reciprocals and then take the reciprocal at the end.",
            "Why: Two equal resistors in parallel always give half the value \u2014 a quick "
            "sanity check.",
            "Common mistake: Forgetting the final \u201Cflip\u201D step. 1/R_T = 1/50. is NOT the "
            "answer \u2014 you must invert to get R_T.",
        ],
    ))

    s6.questions.append(Question(
        qid="6.2",
        type_label="R_T (parallel)",
        difficulty="Straightforward",
        title="Three parallel modules on the ISS",
        prompt=(
            "Three experimental modules on the ISS are wired in parallel across a 12.0-V "
            "regulated bus, with resistances of 60. \u03A9, 30. \u03A9, and 20. \u03A9. "
            "(a) Determine the total resistance seen by the regulator. "
            "(b) Determine the total current the regulator must supply."
        ),
        given=[
            "R\u2081 = 60. \u03A9, R\u2082 = 30. \u03A9, R\u2083 = 20. \u03A9 (parallel)",
            "V = 12.0 V",
        ],
        required=["R_T", "I_total"],
        equation=["1/R_T = 1/R\u2081 + 1/R\u2082 + 1/R\u2083", "I_total = V / R_T"],
        solve=[
            "(a) Common denominator (60):",
            "    1/R_T = 1/60 + 2/60 + 3/60 = 6/60 = 1/10",
            "    R_T = 10. \u03A9",
            "(b) I_total = V / R_T = 12.0 V / 10. \u03A9 = 1.2 A",
        ],
        statement=(
            "The three parallel modules behave like a single 10.-\u03A9 load and draw 1.2 A "
            "in total from the bus."
        ),
        final_answer="R_T = 10. \u03A9; I_total = 1.2 A",
        annotated=[
            "What: Find a common denominator, add the reciprocals, invert.",
            "Why: R_T (= 10 \u03A9) is smaller than even the smallest branch (20 \u03A9). That is "
            "always true for parallel \u2014 useful as a sanity check on diploma exams.",
            "Common mistake: Just averaging the three resistances. Averaging gives 36.7 "
            "\u03A9, which is much larger than 10 \u03A9 \u2014 wrong.",
        ],
    ))

    s6.questions.append(Question(
        qid="6.3",
        type_label="R_T (parallel)",
        difficulty="Challenging",
        title="Five parallel floodlights on a lunar habitat",
        prompt=(
            "Five identical floodlights are mounted on the exterior of a lunar habitat to "
            "illuminate the work area. Each floodlight has a resistance of 96 \u03A9 and they "
            "are all wired in parallel to the habitat\u2019s 120-V power strip. "
            "(a) Determine the total resistance of the floodlight bank. "
            "(b) Determine the total current the power strip must deliver. "
            "(c) Determine the current flowing through any one floodlight. "
            "(d) Verify that the per-light currents add up to your total in (b)."
        ),
        given=[
            "n = 5 floodlights in parallel",
            "R_each = 96 \u03A9",
            "V = 120 V (full source across each light)",
        ],
        required=["R_T", "I_total", "I_each", "verify \u03A3I = I_total"],
        equation=[
            "1/R_T = n / R_each   (when all R are equal)",
            "V = IR (each branch has full source voltage)",
        ],
        solve=[
            "(a) 1/R_T = 5/96  \u2192  R_T = 96 / 5 = 19.2 \u03A9",
            "(b) I_total = V / R_T = 120 V / 19.2 \u03A9 = 6.25 A",
            "(c) Each floodlight has the full 120 V across it, so:",
            "    I_each = V / R_each = 120 V / 96 \u03A9 = 1.25 A",
            "(d) Verification: 5 \u00D7 1.25 A = 6.25 A \u2713",
        ],
        statement=(
            "R_T = 19.2 \u03A9; the power strip must deliver 6.25 A; each floodlight draws "
            "1.25 A, and 5 \u00D7 1.25 A = 6.25 A confirms the result."
        ),
        final_answer="R_T = 19.2 \u03A9; I_total = 6.25 A; I_each = 1.25 A (sum check \u2713)",
        annotated=[
            "What: Use the equal-resistors shortcut R_T = R/n, then check by branch.",
            "Why: Each parallel branch sees the full source voltage. That is fundamentally "
            "different from series, where voltages divide.",
            "Common mistake: Dividing 120 V by R_T to get a single \u201Cbranch\u201D current. "
            "120 / R_T gives the TOTAL current, not the per-branch current.",
        ],
    ))
    sections.append(s6)

    # ===================================================================================
    # TYPE 7: Power
    # ===================================================================================
    s7 = SectionType(
        code="TYPE 7",
        name="Electrical power (P = IV; P = I\u00B2R)",
        equations=["P = IV", "P = I\u00B2R"],
        notes=(
            "Use the form that fits the variables you already have. Both equations give "
            "the same answer if Ohm\u2019s Law is satisfied."
        ),
    )

    s7.questions.append(Question(
        qid="7.1",
        type_label="P = IV",
        difficulty="Straightforward",
        title="Solar array on a science satellite",
        prompt=(
            "A solar array on a small Earth-observation satellite delivers 28.0 V to the "
            "main payload bus, with a steady current draw of 12.0 A while the cameras are "
            "active. "
            "(a) Determine the power delivered to the payload bus. "
            "(b) Later, the satellite enters a low-power mode and the current drops to "
            "8.0 A at the same 28.0 V. Determine the new power."
        ),
        given=[
            "V = 28.0 V",
            "I_active = 12.0 A; I_low = 8.0 A",
        ],
        required=["P_active", "P_low"],
        equation=["P = IV"],
        solve=[
            "(a) P_active = IV = (12.0 A)(28.0 V) = 336 W",
            "(b) P_low    = IV = (8.0 A)(28.0 V)  = 224 W",
        ],
        statement=(
            "The bus draws 336 W in active mode and 224 W in low-power mode."
        ),
        final_answer="P_active = 336 W; P_low = 224 W",
        annotated=[
            "What: Direct multiplication, twice.",
            "Why: P = IV is the most direct power equation when V and I are known.",
            "Common mistake: Forgetting the unit of power. 1 W = 1 V \u00B7 A = 1 J/s.",
        ],
    ))

    s7.questions.append(Question(
        qid="7.2",
        type_label="P = I\u00B2R",
        difficulty="Straightforward",
        title="Heater coil on a Mars rover",
        prompt=(
            "A small heater coil on a Mars rover has a resistance of 4.0 \u03A9. While "
            "operating, it carries a steady current of 2.5 A. "
            "(a) Calculate the power dissipated by the coil using P = I\u00B2R. "
            "(b) Calculate the voltage across the coil using V = IR. "
            "(c) Confirm that P = IV gives the same power as part (a)."
        ),
        given=[
            "R = 4.0 \u03A9",
            "I = 2.5 A",
        ],
        required=["P (using I\u00B2R)", "V (using IR)", "P (using IV)"],
        equation=["P = I\u00B2R", "V = IR", "P = IV"],
        solve=[
            "(a) P = I\u00B2R = (2.5)\u00B2 \u00D7 4.0 = 6.25 \u00D7 4.0 = 25 W",
            "(b) V = IR = (2.5)(4.0) = 10. V",
            "(c) P = IV = (2.5)(10.) = 25 W \u2713  (matches part a)",
        ],
        statement=(
            "The heater coil dissipates 25 W and has 10. V across it. Both forms of the "
            "power equation give the same value, as expected."
        ),
        final_answer="P = 25 W; V = 10. V",
        annotated=[
            "What: Use I\u00B2R first because R and I are given. Then verify with IV.",
            "Why: Diploma-style problems often give you only two of {V, I, R}. Choose the "
            "power form that matches what you have.",
            "Common mistake: Forgetting to square the current. (2.5)\u00B2 = 6.25, not 5.",
        ],
    ))

    s7.questions.append(Question(
        qid="7.3",
        type_label="Power (mixed)",
        difficulty="Challenging",
        title="Solar-powered ISS research module",
        prompt=(
            "A solar-array bus on the ISS supplies a steady 28.0 V to a research module "
            "whose effective resistance is 7.00 \u03A9. "
            "(a) Determine the current the module draws. "
            "(b) Determine the power consumed using P = IV. "
            "(c) Determine the power consumed using P = I\u00B2R, and confirm both forms "
            "agree to within sig figs. "
            "(d) Briefly explain why we still teach two forms of the power equation."
        ),
        given=[
            "V = 28.0 V",
            "R = 7.00 \u03A9",
        ],
        required=["I", "P (two ways)"],
        equation=["V = IR", "P = IV", "P = I\u00B2R"],
        solve=[
            "(a) I = V / R = 28.0 V / 7.00 \u03A9 = 4.00 A",
            "(b) P = IV = (4.00 A)(28.0 V) = 112 W",
            "(c) P = I\u00B2R = (4.00)\u00B2 \u00D7 7.00 = 16.0 \u00D7 7.00 = 112 W \u2713",
            "(d) Two forms exist because real problems can give you any pair of "
            "{V, I, R}. The form you pick avoids extra division and rounding errors.",
        ],
        statement=(
            "The module draws 4.00 A from the 28.0-V bus and consumes 112 W. Both power "
            "formulas give the same result, confirming consistency with Ohm\u2019s Law."
        ),
        final_answer="I = 4.00 A; P = 112 W (both forms agree)",
        annotated=[
            "What: Solve for current first, then use whichever power equation is faster.",
            "Why: P = IV and P = I\u00B2R are mathematically identical once V = IR is true. "
            "Practising both is good prep for diploma multiple choice where the given "
            "data picks the form for you.",
            "Common mistake: Carrying too few sig figs through the chain. Keep one or two "
            "extra digits during work and round at the very end.",
        ],
    ))
    sections.append(s7)

    # ===================================================================================
    # TYPE 8: Energy E = Pt (and kW.h)
    # ===================================================================================
    s8 = SectionType(
        code="TYPE 8",
        name="Electrical energy (E = Pt; 1 kW\u00B7h = 3.60 \u00D7 10\u2076 J)",
        equations=["E = Pt", "1 kW\u00B7h = 3.60 \u00D7 10\u2076 J"],
        notes=(
            "Use to find total energy used by a device. Power must be in watts and time "
            "in seconds for joules; or power in kW and time in hours for kW\u00B7h."
        ),
    )

    s8.questions.append(Question(
        qid="8.1",
        type_label="E = Pt",
        difficulty="Straightforward",
        title="Mars rover heater energy budget",
        prompt=(
            "A 50.-W instrument heater on a Mars rover is run continuously for 4.00 h "
            "during the cold Martian night to keep the electronics within their operating "
            "temperature range. "
            "(a) Determine the energy used by the heater in joules. "
            "(b) Determine the energy used in kilowatt-hours."
        ),
        given=[
            "P = 50. W",
            "t = 4.00 h",
            "1 kW\u00B7h = 3.60 \u00D7 10\u2076 J",
        ],
        required=["E in J", "E in kW\u00B7h"],
        equation=["E = Pt"],
        solve=[
            "(a) Convert time to seconds:",
            "    t = 4.00 h \u00D7 3600 s/h = 1.44 \u00D7 10\u2074 s",
            "    E = Pt = (50. W)(1.44 \u00D7 10\u2074 s) = 7.2 \u00D7 10\u2075 J",
            "(b) In kW\u00B7h directly:",
            "    P = 50. W = 0.050 kW",
            "    E = (0.050 kW)(4.00 h) = 0.20 kW\u00B7h",
            "    Cross-check: 0.20 kW\u00B7h \u00D7 3.60 \u00D7 10\u2076 J/(kW\u00B7h) = 7.2 \u00D7 10\u2075 J \u2713",
        ],
        statement=(
            "Across one 4.00-h Martian night, the heater uses 7.2 \u00D7 10\u2075 J, equivalent "
            "to 0.20 kW\u00B7h."
        ),
        final_answer="E = 7.2\u00D710\u2075 J = 0.20 kW\u00B7h",
        annotated=[
            "What: Pick units before computing. Use seconds for joules, hours for kW\u00B7h.",
            "Why: The kW\u00B7h is just J expressed in larger units; the conversion factor "
            "from the data booklet links them.",
            "Common mistake: Mixing watts with hours and reporting \u201Cwatt-hours\u201D as if it "
            "were joules. 1 W \u00B7 1 h = 3600 J.",
        ],
    ))

    s8.questions.append(Question(
        qid="8.2",
        type_label="E = Pt",
        difficulty="Straightforward",
        title="ISS LED panel lifetime energy",
        prompt=(
            "A 200.-W LED panel inside an ISS module is on for 8.00 h each day to provide "
            "illumination during the working shift. "
            "(a) Determine the energy used per day in kW\u00B7h. "
            "(b) Determine the energy used in one (Earth) year of 365 days, in kW\u00B7h. "
            "(c) Convert your annual answer to joules."
        ),
        given=[
            "P = 200. W = 0.200 kW",
            "t_day = 8.00 h",
            "365 days/year",
            "1 kW\u00B7h = 3.60 \u00D7 10\u2076 J",
        ],
        required=["E_day", "E_year", "E_year in J"],
        equation=["E = Pt"],
        solve=[
            "(a) Per day:",
            "    E_day = (0.200 kW)(8.00 h) = 1.60 kW\u00B7h",
            "(b) Per year:",
            "    E_year = (1.60 kW\u00B7h/day)(365 days) = 584 kW\u00B7h",
            "(c) Convert:",
            "    E_year = 584 \u00D7 3.60 \u00D7 10\u2076 J = 2.10 \u00D7 10\u2079 J",
        ],
        statement=(
            "The LED panel uses 1.60 kW\u00B7h per day and 584 kW\u00B7h per year, equivalent to "
            "2.10 \u00D7 10\u2079 J."
        ),
        final_answer="E_day = 1.60 kW\u00B7h; E_year = 584 kW\u00B7h = 2.10\u00D710\u2079 J",
        annotated=[
            "What: Compute energy per day, scale by 365, then convert to J.",
            "Why: Energy planning for spacecraft works in kW\u00B7h because batteries and "
            "solar arrays are sized in those units.",
            "Common mistake: Multiplying watts by hours and writing the answer in J. "
            "Watt-hours and joules differ by a factor of 3600.",
        ],
    ))

    s8.questions.append(Question(
        qid="8.3",
        type_label="E = Pt",
        difficulty="Challenging",
        title="Standby drain on a satellite ground station",
        prompt=(
            "A satellite ground-station receiver continually draws 8.0 W of stand-by "
            "power, even when the antenna is idle. Electricity costs 12.4 \u00A2/kW\u00B7h. "
            "(a) Determine the energy used by the receiver in one year (365 days), in "
            "kW\u00B7h. "
            "(b) Determine the cost of running this stand-by drain for one year, in "
            "Canadian dollars. "
            "(c) A retrofit reduces the stand-by drain to 2.0 W. Determine the new annual "
            "cost and the savings per year."
        ),
        given=[
            "P\u2081 = 8.0 W = 0.0080 kW",
            "P\u2082 = 2.0 W = 0.0020 kW (after retrofit)",
            "t = 365 \u00D7 24 = 8760 h",
            "rate = 12.4 \u00A2/kW\u00B7h = $0.124/kW\u00B7h",
        ],
        required=[
            "E in kW\u00B7h (year, original)",
            "annual cost (original)",
            "annual cost (retrofit)",
            "savings per year",
        ],
        equation=["E = Pt", "Cost = E \u00D7 rate"],
        solve=[
            "(a) Original energy per year:",
            "    E\u2081 = (0.0080 kW)(8760 h) \u2248 70 kW\u00B7h",
            "(b) Original cost per year:",
            "    Cost\u2081 = 70 kW\u00B7h \u00D7 $0.124/kW\u00B7h \u2248 $8.7",
            "(c) Retrofit energy and cost:",
            "    E\u2082    = (0.0020 kW)(8760 h) \u2248 18 kW\u00B7h",
            "    Cost\u2082 = 18 \u00D7 $0.124 \u2248 $2.2",
            "    Savings = $8.7 \u2212 $2.2 = $6.5 per year",
        ],
        statement=(
            "The original 8.0-W standby drain costs about $8.7/year. After the retrofit "
            "(2.0 W), the cost drops to about $2.2/year, saving roughly $6.5/year per "
            "ground-station receiver."
        ),
        final_answer="E\u2081 \u2248 70 kW\u00B7h; Cost\u2081 \u2248 $8.7; savings \u2248 $6.5/year",
        annotated=[
            "What: Multiply power by total annual hours, then by the rate.",
            "Why: Tiny standby loads multiplied by 8760 h/year add up to surprising "
            "totals \u2014 a key concept on the diploma.",
            "Common mistake: Forgetting to convert cents to dollars (or vice versa). "
            "12.4 \u00A2 = $0.124, not $12.40.",
        ],
    ))
    sections.append(s8)

    # ===================================================================================
    # TYPE 9: Transformers
    # ===================================================================================
    s9 = SectionType(
        code="TYPE 9",
        name="Ideal transformers (N_p/N_s = V_p/V_s = I_s/I_p)",
        equations=[
            "N_p / N_s = V_p / V_s",
            "N_p / N_s = I_s / I_p",
            "V_p / V_s = I_s / I_p",
        ],
        notes=(
            "Step-up transformers raise V (more turns on secondary, less I); step-down "
            "transformers lower V (fewer turns on secondary, more I). Ideal transformers "
            "conserve power: V_p I_p = V_s I_s."
        ),
    )

    s9.questions.append(Question(
        qid="9.1",
        type_label="Transformer",
        difficulty="Straightforward",
        title="Mars-colony habitat power transmission",
        prompt=(
            "A future Mars colony habitat uses a step-up transformer to take a 240-V solar "
            "array input and raise it to 12 000 V for long-distance power lines linking "
            "neighbouring habitats. The primary coil has 125 turns. "
            "(a) Determine whether the transformer is step-up or step-down, and the number "
            "of turns required on the secondary coil. "
            "(b) If the primary draws 25.0 A, determine the current in the secondary."
        ),
        given=[
            "V_p = 240 V; V_s = 12 000 V",
            "N_p = 125 turns",
            "I_p = 25.0 A",
        ],
        required=["type (step-up/step-down)", "N_s", "I_s"],
        equation=[
            "N_p / N_s = V_p / V_s  \u2192  N_s = N_p \u00D7 (V_s / V_p)",
            "I_s / I_p = V_p / V_s  \u2192  I_s = I_p \u00D7 (V_p / V_s)",
        ],
        solve=[
            "(a) V_s > V_p, so this is a step-up transformer.",
            "    N_s = 125 \u00D7 (12 000 / 240) = 125 \u00D7 50.0 = 6250 turns",
            "(b) I_s = 25.0 \u00D7 (240 / 12 000) = 25.0 \u00D7 0.0200 = 0.500 A",
        ],
        statement=(
            "It is a step-up transformer: the secondary needs 6250 turns and the secondary "
            "current drops to 0.500 A while the voltage rises by a factor of 50."
        ),
        final_answer="Step-up; N_s = 6250 turns; I_s = 0.500 A",
        annotated=[
            "What: Use the turn-voltage ratio to find N_s, then the current ratio to find "
            "I_s.",
            "Why: Power is conserved in an ideal transformer: V_p I_p = V_s I_s. Higher V "
            "automatically means lower I.",
            "Common mistake: Flipping which side gets more turns. Memorise: more turns "
            "on the side with higher voltage.",
        ],
    ))

    s9.questions.append(Question(
        qid="9.2",
        type_label="Transformer",
        difficulty="Straightforward",
        title="Spacecraft step-down for instrument power",
        prompt=(
            "A spacecraft instrument is rated for 5.00 V, but the vehicle\u2019s main power bus "
            "supplies 28.0 V. A step-down transformer is used. The primary coil has 560 "
            "turns and draws 0.500 A from the bus. "
            "(a) Determine the number of turns required on the secondary coil. "
            "(b) Determine the current that the secondary delivers to the instrument."
        ),
        given=[
            "V_p = 28.0 V; V_s = 5.00 V",
            "N_p = 560 turns",
            "I_p = 0.500 A",
        ],
        required=["N_s", "I_s"],
        equation=[
            "N_p / N_s = V_p / V_s",
            "N_p / N_s = I_s / I_p",
        ],
        solve=[
            "(a) N_s = N_p \u00D7 (V_s / V_p) = 560 \u00D7 (5.00 / 28.0) = 560 \u00D7 0.1786",
            "    N_s = 100. turns",
            "(b) I_s = I_p \u00D7 (N_p / N_s) = 0.500 \u00D7 (560 / 100.) = 0.500 \u00D7 5.60",
            "    I_s = 2.80 A",
        ],
        statement=(
            "The secondary coil needs 100. turns and delivers 2.80 A to the instrument at "
            "5.00 V."
        ),
        final_answer="N_s = 100. turns; I_s = 2.80 A",
        annotated=[
            "What: Apply the turn-voltage ratio for N_s, then the turn-current ratio for "
            "I_s.",
            "Why: Stepping voltage DOWN steps current UP, so the instrument sees more A "
            "than the bus.",
            "Common mistake: Using V_p for V_s in the ratio. Always label which coil is "
            "primary and which is secondary in your diagram before you substitute.",
        ],
    ))

    s9.questions.append(Question(
        qid="9.3",
        type_label="Transformer",
        difficulty="Challenging",
        title="Step-up for a deep-space transmission line",
        prompt=(
            "A future Mars-orbit communications hub uses a step-up transformer to send "
            "power along a 7.20 \u00D7 10\u00B3 V transmission cable that supplies an array of "
            "antennas. The primary side runs at 240 V from a fuel-cell stack and has 80 "
            "turns. The fuel cells deliver 1.00 \u00D7 10\u00B2 A to the primary. "
            "(a) Determine the number of turns on the secondary coil. "
            "(b) Determine the current carried in the high-voltage transmission cable. "
            "(c) Verify that the input power equals the output power for an ideal "
            "transformer."
        ),
        given=[
            "V_p = 240 V; V_s = 7.20 \u00D7 10\u00B3 V",
            "N_p = 80 turns",
            "I_p = 1.00 \u00D7 10\u00B2 A",
        ],
        required=["N_s", "I_s", "P_in vs. P_out check"],
        equation=[
            "N_p / N_s = V_p / V_s",
            "I_s / I_p = V_p / V_s",
            "P = IV (for the verification)",
        ],
        solve=[
            "(a) N_s = N_p \u00D7 (V_s / V_p) = 80 \u00D7 (7200 / 240) = 80 \u00D7 30 = 2400 turns",
            "(b) I_s = I_p \u00D7 (V_p / V_s) = 100 \u00D7 (240 / 7200) = 100 \u00D7 0.03333",
            "    I_s = 3.33 A",
            "(c) Power check:",
            "    P_in  = V_p I_p = (240)(100.) = 2.40 \u00D7 10\u2074 W",
            "    P_out = V_s I_s = (7200)(3.33) = 2.40 \u00D7 10\u2074 W \u2713",
        ],
        statement=(
            "Secondary has 2400 turns and carries 3.33 A; both the input and output powers "
            "equal 2.40 \u00D7 10\u2074 W, confirming ideal-transformer behaviour."
        ),
        final_answer="N_s = 2400 turns; I_s = 3.33 A; P_in = P_out = 2.40\u00D710\u2074 W",
        annotated=[
            "What: Two ratio calculations, then a power-conservation check.",
            "Why: Power-loss reduction is the whole point of high-voltage transmission \u2014 "
            "lower I means I\u00B2R losses fall sharply along long cables.",
            "Common mistake: Applying P = IV with mismatched primary/secondary numbers. "
            "Keep your subscripts straight: primary together, secondary together.",
        ],
    ))
    sections.append(s9)

    # ===================================================================================
    # TYPE 10: Wave equation
    # ===================================================================================
    s10 = SectionType(
        code="TYPE 10",
        name="Universal wave equation (v = f\u03BB; for EMR c = f\u03BB)",
        equations=["v = f\u03BB", "c = f\u03BB", "c = 3.00 \u00D7 10\u2078 m/s"],
        notes=(
            "Use v = f\u03BB for any wave; for electromagnetic radiation in vacuum/air, "
            "v = c = 3.00 \u00D7 10\u2078 m/s. Frequency in Hz, wavelength in m."
        ),
    )

    s10.questions.append(Question(
        qid="10.1",
        type_label="c = f\u03BB",
        difficulty="Straightforward",
        title="Hubble observes visible-band EMR",
        prompt=(
            "The Hubble Space Telescope detects electromagnetic radiation from a distant "
            "galaxy with a frequency of 6.00 \u00D7 10\u00B9\u2074 Hz. "
            "(a) Determine the wavelength of this radiation. "
            "(b) Use the EMR spectrum diagram on data booklet page 3 to identify which "
            "region of the EM spectrum this wavelength belongs to."
        ),
        given=[
            "f = 6.00 \u00D7 10\u00B9\u2074 Hz",
            "c = 3.00 \u00D7 10\u2078 m/s",
        ],
        required=["\u03BB", "spectrum region"],
        equation=["c = f\u03BB  \u2192  \u03BB = c / f"],
        solve=[
            "(a) \u03BB = c / f = (3.00 \u00D7 10\u2078) / (6.00 \u00D7 10\u00B9\u2074) = 5.00 \u00D7 10\u207B\u2077 m",
            "(b) 5.00 \u00D7 10\u207B\u2077 m = 500 nm \u2192 visible light (green region of ROYGBIV).",
        ],
        statement=(
            "The radiation has a wavelength of \u03BB = 5.00 \u00D7 10\u207B\u2077 m (500 nm), placing it "
            "in the visible-light portion of the EM spectrum."
        ),
        final_answer="\u03BB = 5.00 \u00D7 10\u207B\u2077 m (500 nm) \u2014 visible light",
        annotated=[
            "What: Rearrange c = f\u03BB to solve for \u03BB, then look up the band.",
            "Why: c is fixed in vacuum/air, so \u03BB and f are inverses. Higher f \u2192 shorter "
            "\u03BB.",
            "Common mistake: Forgetting that the data booklet spectrum diagram is in metres "
            "(bottom scale). 5 \u00D7 10\u207B\u2077 m is between 10\u207B\u2076 m and 10\u207B\u2078 m \u2014 the visible "
            "band.",
        ],
    ))

    s10.questions.append(Question(
        qid="10.2",
        type_label="c = f\u03BB",
        difficulty="Straightforward",
        title="Cassini downlink to Earth",
        prompt=(
            "The Cassini probe sent data back to Earth using a radio downlink with a "
            "wavelength of 0.130 m in vacuum. "
            "(a) Determine the frequency of the downlink signal. "
            "(b) If a future mission uses radio with twice the wavelength of Cassini\u2019s "
            "downlink (i.e. \u03BB = 0.260 m), determine its frequency."
        ),
        given=[
            "\u03BB\u2081 = 0.130 m",
            "\u03BB\u2082 = 0.260 m",
            "c = 3.00 \u00D7 10\u2078 m/s",
        ],
        required=["f\u2081", "f\u2082"],
        equation=["c = f\u03BB  \u2192  f = c / \u03BB"],
        solve=[
            "(a) f\u2081 = c / \u03BB\u2081 = (3.00 \u00D7 10\u2078) / 0.130",
            "    f\u2081 = 2.31 \u00D7 10\u2079 Hz   (\u2248 2.31 GHz, S/X-band radio)",
            "(b) f\u2082 = c / \u03BB\u2082 = (3.00 \u00D7 10\u2078) / 0.260",
            "    f\u2082 = 1.15 \u00D7 10\u2079 Hz   (half of f\u2081 \u2014 expected, since \u03BB doubled)",
        ],
        statement=(
            "Cassini\u2019s downlink used f\u2081 \u2248 2.31 \u00D7 10\u2079 Hz. Doubling the wavelength halves "
            "the frequency to f\u2082 \u2248 1.15 \u00D7 10\u2079 Hz."
        ),
        final_answer="f\u2081 \u2248 2.31\u00D710\u2079 Hz; f\u2082 \u2248 1.15\u00D710\u2079 Hz",
        annotated=[
            "What: Two direct uses of f = c/\u03BB.",
            "Why: Demonstrates the inverse relationship between f and \u03BB at fixed c.",
            "Common mistake: Switching f and \u03BB in the equation. Cross-check: the answer "
            "for radio should land between 10\u2078 and 10\u00B9\u2070 Hz on the data-booklet diagram.",
        ],
    ))

    s10.questions.append(Question(
        qid="10.3",
        type_label="c = f\u03BB",
        difficulty="Challenging",
        title="Red-shift of H-alpha from a receding galaxy",
        prompt=(
            "Astronomers detect the H-alpha emission line, which has a rest wavelength of "
            "656 nm in the laboratory. From a distant galaxy, the same line is observed "
            "shifted to 660 nm. "
            "(a) Determine the lab (\u201Cunshifted\u201D) frequency f. "
            "(b) Determine the observed (red-shifted) frequency f\u2032. "
            "(c) Determine the change in frequency \u0394f = f \u2212 f\u2032 and state, with reasoning, "
            "whether the galaxy is moving toward or away from Earth."
        ),
        given=[
            "\u03BB_lab = 656 nm = 6.56 \u00D7 10\u207B\u2077 m",
            "\u03BB_obs = 660 nm = 6.60 \u00D7 10\u207B\u2077 m",
            "c = 3.00 \u00D7 10\u2078 m/s",
        ],
        required=["f (lab)", "f\u2032 (observed)", "\u0394f and direction of motion"],
        equation=["c = f\u03BB  \u2192  f = c / \u03BB"],
        solve=[
            "(a) Lab frequency:",
            "    f  = c / \u03BB_lab = (3.00 \u00D7 10\u2078) / (6.56 \u00D7 10\u207B\u2077)",
            "    f  = 4.57 \u00D7 10\u00B9\u2074 Hz",
            "(b) Observed (shifted) frequency:",
            "    f\u2032 = c / \u03BB_obs = (3.00 \u00D7 10\u2078) / (6.60 \u00D7 10\u207B\u2077)",
            "    f\u2032 = 4.55 \u00D7 10\u00B9\u2074 Hz",
            "(c) Carry one extra digit so the difference is not lost to rounding:",
            "    f  = 4.573 \u00D7 10\u00B9\u2074 Hz",
            "    f\u2032 = 4.545 \u00D7 10\u00B9\u2074 Hz",
            "    \u0394f = f \u2212 f\u2032 = 0.028 \u00D7 10\u00B9\u2074 Hz \u2248 2.8 \u00D7 10\u00B9\u00B2 Hz",
            "    Lower observed frequency \u2192 RED shift \u2192 the source is moving AWAY from",
            "    Earth, per the Doppler effect.",
        ],
        statement=(
            "The lab frequency is 4.57 \u00D7 10\u00B9\u2074 Hz. The observed frequency drops to "
            "4.55 \u00D7 10\u00B9\u2074 Hz \u2014 a decrease of about 2.8 \u00D7 10\u00B9\u00B2 Hz. The lower observed "
            "frequency is a red shift, indicating the galaxy is receding from Earth."
        ),
        final_answer=(
            "f = 4.57\u00D710\u00B9\u2074 Hz; f\u2032 = 4.55\u00D710\u00B9\u2074 Hz; \u0394f \u2248 2.8\u00D710\u00B9\u00B2 Hz \u2014 red-shifted (receding)"
        ),
        annotated=[
            "What: Compute two frequencies, then interpret the shift physically.",
            "Why: Red shift is one of the strongest pieces of evidence that the universe "
            "is expanding \u2014 a key Unit C concept.",
            "Common mistake: Reporting \u0394f as a tiny rounding artifact. Carry an extra "
            "sig fig during the subtraction so that \u0394f is meaningful.",
        ],
    ))
    sections.append(s10)

    return sections


# --------------------------------------------------------------------------------------
# DOCX rendering
# --------------------------------------------------------------------------------------


def _set_base_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"


def _add_para(doc: Document, text: str = "", *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic


def _add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _add_label_block(doc: Document, label: str, lines: Iterable[str]) -> None:
    """Bold label paragraph then indented body lines."""
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    for line in lines:
        body = doc.add_paragraph(line)
        body.paragraph_format.left_indent = Cm(0.6)


def render_docx(sections: list[SectionType]) -> Path:
    doc = Document()
    _set_base_style(doc)

    # Title
    _add_heading(doc, "Science 30 \u2014 Unit C: Electromagnetic Energy", level=0)
    _add_para(doc, "Diploma-level Review and Test-Prep Question Set", bold=True)
    _add_para(doc, "Student Booklet WITH Worked GRAS(S) Solutions", italic=True)

    # Summary of inputs used
    _add_heading(doc, "Summary of Inputs Used", level=1)
    _add_bullets(doc, [
        "Course: Science 30 (Alberta program of studies)",
        "Unit: Unit C \u2014 Electromagnetic Energy",
        "Lessons covered: Lesson 1 Fields, Lesson 2 Electric/Magnetic Fields, "
        "Lesson 3 Motors and Generators, Lesson 4 Circuits, Lesson 5 Transmitting "
        "Electrical Energy, Lesson 6 EM Radiation, Lesson 7 Properties of Light, "
        "Lesson 8 Astronomy",
        "Source PPT: \u201CScience 30 - Unit C - Electromagnetic Energy.pptx\u201D",
        "Review sheet: \u201CBONUS - Unit 3 Review Sheet.docx\u201D",
        "Allowed data: Alberta Science 30 data booklet pp. 2\u20133 only "
        "(gravitational/electric fields, astronomy data, electricity formulas, "
        "wave formulas, EM spectrum)",
        "Question count: 30 (10 equation types \u00D7 3 questions: 2 straightforward + 1 challenging)",
        "Format: every question is a multi-step word problem",
        "Method: GRAS(S) \u2014 Given, Required, Applicable equation, Substitute and solve, Statement",
        "Sig figs: final answers carry the smallest sig-fig count of the given data",
        "Themes: space-themed contexts (Rosetta/Philae, Mars Perseverance, ISS, Hubble, "
        "Cassini, Apollo Moon dust, lunar habitat, deep-space probes, asteroid Bennu, "
        "Mars colony power transmission)",
    ])

    # How to use
    _add_heading(doc, "How to Use This Booklet", level=1)
    _add_para(doc,
        "Each question below is followed immediately by its full GRAS(S) solution and "
        "annotated thinking. Try the question on a separate sheet first, then check the "
        "worked solution to find any step where your reasoning diverged. Pay close "
        "attention to the \u201CCommon mistake\u201D notes \u2014 these are the highest-leverage "
        "diploma-exam pitfalls.")
    _add_para(doc,
        "GRAS(S) reminder: G \u2014 Given. R \u2014 Required. A \u2014 Applicable equation. "
        "S \u2014 Substitute and solve. S \u2014 Statement with units and significant figures.")

    # Sections
    for sec in sections:
        _add_heading(doc, f"{sec.code}: {sec.name}", level=1)
        _add_label_block(doc, "Equation(s) used:", sec.equations)
        _add_para(doc, f"When to use: {sec.notes}", italic=True)

        for q in sec.questions:
            _add_heading(doc, f"Question {q.qid} ({q.difficulty}) \u2014 {q.title}", level=2)
            _add_para(doc, q.prompt)

            _add_label_block(doc, "G \u2014 Given:", q.given)
            _add_label_block(doc, "R \u2014 Required:", q.required)
            _add_label_block(doc, "A \u2014 Applicable equation(s):", q.equation)
            _add_label_block(doc, "S \u2014 Substitute and solve:", q.solve)
            _add_label_block(doc, "S \u2014 Statement (units + sig figs):", [q.statement])

            p = doc.add_paragraph()
            p.add_run("Final answer: ").bold = True
            p.add_run(q.final_answer)

            _add_label_block(doc, "Annotated thinking (what / why / common mistake):", q.annotated)

    # Audit
    _add_heading(doc, "Audit / Traceability Notes", level=1)
    _add_bullets(doc, [
        "Outcome mapping by equation type: see TYPE 1\u201310 headings; each maps to one "
        "data-booklet equation.",
        "PPT lessons covered (calculation-based content): "
        "Lesson 1 (gravitational fields) \u2192 TYPES 1\u20132; "
        "Lesson 2 (electric fields) \u2192 TYPE 3; "
        "Lesson 4 (circuits, Ohm\u2019s Law, series, parallel) \u2192 TYPES 4\u20136; "
        "Lesson 5 (power, energy, transformers) \u2192 TYPES 7\u20139; "
        "Lesson 6 (EM radiation) \u2192 TYPE 10; "
        "Lesson 8 (Doppler / red-shift) \u2192 reinforced in Q10.3.",
        "Review-sheet items addressed quantitatively: gravitational fields, electric "
        "fields, circuits (series/parallel/Voltage/Current/Resistance), formulas \u00D7 6, "
        "transformers, EM waves and frequency/wavelength, red-shift application.",
        "Review-sheet items left for separate qualitative review: magnetic fields, "
        "motors vs. generators, AC vs. DC, properties of light (reflection, refraction, "
        "polarization, diffraction), spectroscopy diagrams, evolution of stars. These "
        "do not have data-booklet equations and so were excluded from this build by "
        "scope.",
        "All numerical results were independently re-verified against the constants on "
        "data booklet p. 2\u20133.",
        "Difficulty distribution per type: 2 \u00D7 Straightforward + 1 \u00D7 Challenging.",
    ])

    # Uncertainty
    _add_heading(doc, "Uncertainty / Review Flags", level=1)
    _add_bullets(doc, [
        "Comet 67P touchdown surface gravity is treated as 1.0 \u00D7 10\u207B\u00B3 N/kg in Q1.1; "
        "real values vary across the comet\u2019s irregular surface (10\u207B\u2074 to 10\u207B\u00B3 N/kg). "
        "The number is approximate but pedagogically representative.",
        "Asteroid Bennu surface g (7.8 \u00D7 10\u207B\u2075 N/kg) used in Q1.3 is an order-of-"
        "magnitude estimate; treat as approximate context, not as precise data.",
        "ISS altitude in Q2.1 is taken as 4.20 \u00D7 10\u2075 m (\u2248 420 km). The ISS actually "
        "varies between roughly 400\u2013420 km and is reboosted periodically; the value is "
        "good to about 5 %.",
        "Q10.3 uses a simplified red-shift framing (compare lab vs. observed \u03BB) without "
        "invoking the full relativistic Doppler equation, which is appropriate at the "
        "Science 30 level per the data booklet.",
        "Question 9.1 numbers are intentionally identical to a PPT example (240 V \u2192 "
        "12 000 V, 125 turns, 25.0 A) to provide a direct mirror exercise; teacher may "
        "wish to swap numbers if assigning this as a graded assessment.",
    ])

    # Suggested next step
    _add_heading(doc, "Suggested Next Step", level=1)
    _add_para(doc,
        "Recommended classroom use: assign as a unit-test review pack one class day "
        "before the Unit C test, with students attempting all 30 questions on a "
        "separate sheet first and using the worked GRAS(S) solutions only after a "
        "good-faith attempt. Pair with a separate qualitative review sheet (or the "
        "BONUS sheet itself) to cover magnetic fields, AC/DC, optics, spectroscopy, "
        "Doppler, and stellar evolution items that do not have data-booklet equations. "
        "Optional follow-up: a short formative quiz containing one question per "
        "equation type \u2014 mirror these contexts with different numbers.")

    doc.save(DOCX_PATH)
    return DOCX_PATH


# --------------------------------------------------------------------------------------
# Markdown summary (template-aligned)
# --------------------------------------------------------------------------------------


def render_markdown_summary(sections: list[SectionType]) -> Path:
    lines: list[str] = []
    lines.append("# Science 30 \u2014 Unit C: Electromagnetic Energy")
    lines.append("## Review and Test Prep \u2014 Audit / Summary Companion")
    lines.append("")
    lines.append("> Companion to `sci30-unitc-emr-review-student.docx`. The DOCX is the "
                 "primary student-facing deliverable (questions and full GRAS(S) "
                 "solutions inline). This file is the audit/traceability companion in "
                 "the format of `templates/review-test-prep-template.md`.")
    lines.append("")
    lines.append("## Summary of Inputs Used")
    lines.append("- Course: Science 30 (Alberta)")
    lines.append("- Unit/topic: Unit C \u2014 Electromagnetic Energy")
    lines.append("- Outcomes/standards targeted: gravitational and electric fields; "
                 "circuits (Ohm\u2019s Law, series, parallel); power and energy; "
                 "transformers and high-voltage transmission; universal wave equation "
                 "and EM radiation; red-shift application of c = f\u03BB.")
    lines.append("- Constraints: 10 equation types \u00D7 3 questions = 30 questions total; "
                 "2 straightforward + 1 challenging per type; every question multi-step "
                 "and word-problem; Alberta Science 30 data booklet pp. 2\u20133 only; "
                 "calculator allowed; sig figs governed by least-precise given.")
    lines.append("")
    lines.append("## Question Set")
    for sec in sections:
        lines.append(f"### {sec.code}: {sec.name}")
        lines.append(f"- Equation(s): {'; '.join(sec.equations)}")
        for q in sec.questions:
            lines.append(f"- **Q{q.qid}** ({q.difficulty}) \u2014 {q.title}: {q.prompt}")
        lines.append("")
    lines.append("## Fully Worked Student-Facing Key")
    lines.append("Full GRAS(S) keys with annotated thinking are inline in the DOCX. "
                 "Every question includes G, R, A, S (substitute & solve), S (statement "
                 "with units and sig figs), final answer, and a what/why/common-mistake "
                 "block.")
    lines.append("")
    lines.append("## Audit / Traceability Notes")
    lines.append("- Outcome mapping by equation type:")
    for sec in sections:
        eqstr = "; ".join(sec.equations)
        lines.append(f"  - {sec.code} ({sec.name}): {eqstr}")
    lines.append("- Skill coverage by question difficulty: each type has 2 straightforward "
                 "(direct application + 1 instructional twist such as a unit conversion "
                 "or a small comparison) and 1 challenging question (multi-equation, "
                 "multi-part, with a verification or interpretation step).")
    lines.append("- Assumptions made: see Uncertainty/Review Flags below.")
    lines.append("")
    lines.append("## Uncertainty or Review Flags")
    lines.append("- Q1.1: 67P surface g \u2248 1.0 \u00D7 10\u207B\u00B3 N/kg is approximate (varies across "
                 "the comet\u2019s irregular surface).")
    lines.append("- Q1.3: Bennu surface g \u2248 7.8 \u00D7 10\u207B\u2075 N/kg is order-of-magnitude.")
    lines.append("- Q2.1: ISS altitude treated as 4.20 \u00D7 10\u2075 m; actual altitude drifts "
                 "in the 400\u2013420 km range with periodic reboosts.")
    lines.append("- Q9.1: Numbers intentionally mirror a PPT worked example (240 V \u2192 "
                 "12 000 V, 125 turns, 25.0 A). Swap numbers if used in a graded "
                 "assessment.")
    lines.append("- Q10.3: Doppler red-shift treated qualitatively (compare \u03BB and f at "
                 "lab vs. observed). Full relativistic Doppler is out of scope for "
                 "Science 30.")
    lines.append("- Items deliberately out of scope (no data-booklet equation): magnetic "
                 "fields qualitative review, motors vs. generators, AC vs. DC, optics "
                 "(reflection / refraction / polarization / diffraction), spectroscopy "
                 "diagrams, evolution of stars.")
    lines.append("")
    lines.append("## Suggested Next Step")
    lines.append("- Use the day before the Unit C test as a guided-review session: "
                 "students attempt all 30 questions independently on a separate sheet, "
                 "then check against the inline GRAS(S) keys. Pair with the BONUS Unit 3 "
                 "review sheet to cover the qualitative items not addressed here.")
    lines.append("")
    lines.append("## Standard Agent Contract")
    lines.append("- **In scope:** equation-based, multi-step word problems aligned to the "
                 "10 equation types in the Alberta Science 30 data booklet pp. 2\u20133 that "
                 "are taught in Unit C. Full GRAS(S) keys with annotated thinking. "
                 "DOCX deliverable. Markdown audit companion.")
    lines.append("- **Out of scope:** qualitative review-sheet items with no data-booklet "
                 "equation (magnetic fields, motors/generators, AC/DC, optics, spectroscopy, "
                 "stellar evolution); diagrams (field-line drawings, circuit schematics); "
                 "lab procedure questions; Part A multiple-choice formatting.")
    lines.append("- **Completed:** 30 multi-step questions (10 types \u00D7 3) with worked "
                 "GRAS(S) solutions, sig-fig-aware final statements, annotated thinking "
                 "(what/why/common mistake), audit map, uncertainty flags, suggested "
                 "next step. DOCX and MD generated.")
    lines.append("- **Not completed:** separate teacher-only key (one combined "
                 "student+key DOCX was requested); diagrams; multiple-choice or numerical-"
                 "response formatting; qualitative review questions.")
    lines.append("- **Assumptions and decisions:** combined power formulas (P=IV and "
                 "P=I\u00B2R) into a single \u201Cpower\u201D type with three questions; combined "
                 "v=f\u03BB and c=f\u03BB into a single wave-equation type; used Earth-Moon, Mars, "
                 "ISS, comet 67P, Bennu, and lunar-habitat scenarios as the space-themed "
                 "contexts.")
    lines.append("- **Risks or uncertainty:** approximate physical values flagged in "
                 "Uncertainty section; numbers in Q9.1 mirror a PPT example.")
    lines.append("")

    MD_PATH.write_text("\n".join(lines), encoding="utf-8")
    return MD_PATH


# --------------------------------------------------------------------------------------
# Entry point
# --------------------------------------------------------------------------------------


if __name__ == "__main__":
    sections = build_sections()
    n_questions = sum(len(s.questions) for s in sections)
    print(f"Built {len(sections)} equation-type sections / {n_questions} questions.")

    docx_path = render_docx(sections)
    print(f"DOCX -> {docx_path}")

    md_path = render_markdown_summary(sections)
    print(f"MD   -> {md_path}")
